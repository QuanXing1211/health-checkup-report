"""
html_to_word_export 单元测试
============================
用例覆盖：配置加载、DOM 提取、DOM → docx 映射、表格合并单元格、
base64 图片、截图占位、风险卡 / KPI / 标签 / 评级徽章、skip 节点剔除。

运行：
    python -m unittest html_to_word.tests.test_html_to_word_export -v
"""

import io
import os
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import MagicMock, patch

# 注入项目根到 sys.path
PROJECT_ROOT = Path(__file__).resolve().parents[2]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from bs4 import BeautifulSoup

from html_to_word.html_to_word_export import HtmlToWordExporter, _Container


FIXTURE_DIR = Path(__file__).resolve().parent / "fixtures"
MINI_HTML = FIXTURE_DIR / "mini-report.html"


def _make_exporter(tmpdir, html_path=None, output_name="out.docx"):
    """构造一个不渲染的 exporter，用于直接测 assemble_dom。

    默认关闭 COM 更新域（避免单测启动真实 Word 进程）；
    TestUpdateFieldsViaWord 用例按需在用例内重新开启。
    """
    html_path = html_path or MINI_HTML
    output_path = Path(tmpdir) / output_name
    exp = HtmlToWordExporter(
        input_path=html_path,
        output_path=output_path,
        config_path=None,
    )
    exp.config["update_fields_after_save"] = {
        "enabled": False, "backend": "none", "fail_silently": True,
    }
    return exp


class _TestBase(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.mkdtemp(prefix="h2w_test_")
    def tearDown(self):
        import shutil
        shutil.rmtree(self.tmpdir, ignore_errors=True)


class TestLoadConfig(_TestBase):
    def test_load_config_default_when_no_path(self):
        exp = _make_exporter(self.tmpdir)
        cfg = exp.load_config()
        self.assertEqual(cfg["preview_mode"], "a4-landscape")
        self.assertIn("style_map", cfg)
        self.assertEqual(cfg["docx"]["page_width_mm"], 297)

    def test_load_config_custom_yaml_overrides(self):
        custom = Path(self.tmpdir) / "custom.yaml"
        custom.write_text(
            "preview_mode: a4-portrait\n"
            "docx:\n  page_width_mm: 210\n  page_height_mm: 297\n",
            encoding="utf-8",
        )
        exp = _make_exporter(self.tmpdir)
        exp.config_path = custom
        cfg = exp.load_config()
        self.assertEqual(cfg["preview_mode"], "a4-portrait")
        self.assertEqual(cfg["docx"]["page_width_mm"], 210)
        self.assertEqual(cfg["docx"]["page_height_mm"], 297)
        # 未覆盖的字段保留默认
        self.assertEqual(cfg["docx"]["margin_mm"], 20)

    def test_preview_mode_override_in_constructor(self):
        exp = _make_exporter(self.tmpdir)
        exp.preview_mode = None  # 不通过构造器覆盖
        # 构造时传 preview_mode
        exp2 = HtmlToWordExporter(
            input_path=MINI_HTML,
            output_path=Path(self.tmpdir) / "x.docx",
            preview_mode="a4-portrait",
        )
        self.assertEqual(exp2.config["preview_mode"], "a4-portrait")


class TestExtractDom(_TestBase):
    def test_extract_dom_returns_root(self):
        exp = _make_exporter(self.tmpdir)
        # 直接构造 mock page，content 返回 fixture
        html = MINI_HTML.read_text(encoding="utf-8")
        page = MagicMock()
        page.content.return_value = html
        root = exp.extract_dom(page)
        self.assertEqual(root.name, "div")
        self.assertEqual(root.get("id"), "sr-page-source")

    def test_extract_dom_skips_nav_and_toc(self):
        exp = _make_exporter(self.tmpdir)
        html = MINI_HTML.read_text(encoding="utf-8")
        page = MagicMock()
        page.content.return_value = html
        root = exp.extract_dom(page)
        # sr-nav-wrap 应被剔除
        nav = root.find(class_="sr-nav-wrap")
        toc = root.find(class_="sr-toc-page")
        self.assertIsNone(nav)
        self.assertIsNone(toc)
        # 正文应保留
        self.assertIsNotNone(root.find("h1"))

    def test_extract_dom_falls_back_to_body_when_no_root(self):
        exp = _make_exporter(self.tmpdir)
        exp.config["dom_root"] = "#does-not-exist"
        html = MINI_HTML.read_text(encoding="utf-8")
        page = MagicMock()
        page.content.return_value = html
        root = exp.extract_dom(page)
        # 回退到 body
        self.assertIsNotNone(root)
        self.assertIn(root.name, ("div", "body"))


class TestAssembleDocx(_TestBase):
    def _assemble(self, html, snapshot_map=None, config_overrides=None):
        exp = _make_exporter(self.tmpdir)
        if config_overrides:
            for k, v in config_overrides.items():
                exp.config[k] = v
        exp._snapshot_map = snapshot_map or {}
        root = BeautifulSoup(html, "lxml").find(id="sr-page-source") or BeautifulSoup(html, "lxml").body
        out = Path(self.tmpdir) / "out.docx"
        exp.assemble_docx(root, exp._snapshot_map, out)
        from docx import Document
        return Document(str(out))

    def test_headings_mapped_to_levels(self):
        html = '<div id="sr-page-source"><h1>H1</h1><h2>H2</h2><h3>H3</h3><h4>H4</h4><h5>H5</h5><h6>H6</h6></div>'
        doc = self._assemble(html)
        levels = [p.style.name for p in doc.paragraphs if p.text]
        self.assertIn("Heading 1", levels)
        self.assertIn("Heading 2", levels)
        self.assertIn("Heading 3", levels)
        self.assertIn("Heading 4", levels)
        self.assertIn("Heading 5", levels)
        self.assertIn("Heading 6", levels)

    def test_paragraph_text(self):
        html = '<div id="sr-page-source"><p>正文段落 A</p><p>正文段落 B</p></div>'
        doc = self._assemble(html)
        texts = [p.text for p in doc.paragraphs if p.text]
        self.assertIn("正文段落 A", texts)
        self.assertIn("正文段落 B", texts)

    def test_table_with_colspan(self):
        html = '''<div id="sr-page-source"><table>
            <tr><th>序号</th><th>设备</th><th>状态</th></tr>
            <tr><td>1</td><td colspan="2">防火墙合并</td></tr>
            <tr><td>2</td><td>IDS</td><td>正常</td></tr>
        </table></div>'''
        doc = self._assemble(html)
        self.assertEqual(len(doc.tables), 1)
        t = doc.tables[0]
        # 3 行 3 列
        self.assertEqual(len(t.rows), 3)
        # 第 2 行第 2 个单元格文本为合并内容
        self.assertEqual(t.cell(1, 1).text, "防火墙合并")

    def test_base64_image_inserted(self):
        # 1x1 透明 PNG
        b64 = "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNkYAAAAAYAAjCB0C8AAAAASUVORK5CYII="
        html = f'<div id="sr-page-source"><img src="data:image/png;base64,{b64}"/></div>'
        doc = self._assemble(html)
        # 至少应包含 inline image
        from docx.document import Document as _DocCls
        # 统计 inline shapes
        self.assertGreaterEqual(len(doc.inline_shapes), 1)

    def test_snapshot_placeholder_inserted_from_map(self):
        # 准备一张临时图片
        from PIL import Image  # noqa: F401  # 仅当 PIL 可用时用，否则跳过
        try:
            from PIL import Image
        except ImportError:
            self.skipTest("PIL not installed")

        img_path = Path(self.tmpdir) / "snap.png"
        Image.new("RGB", (10, 10), "white").save(str(img_path))

        html = '<div id="sr-page-source"><img data-snapshot="snap-001" alt="chart"/></div>'
        doc = self._assemble(html, snapshot_map={"snap-001": str(img_path)})
        self.assertGreaterEqual(len(doc.inline_shapes), 1)

    def test_risk_card_table_created(self):
        html = '''<div id="sr-page-source"><div class="sr-risk-card">
            <h6>问题描述</h6><p>测试风险</p>
            <h6>影响</h6><p>影响说明</p>
        </div></div>'''
        exp = _make_exporter(self.tmpdir)
        exp._snapshot_map = {}
        root = BeautifulSoup(html, "lxml").find(id="sr-page-source")
        out = Path(self.tmpdir) / "rc.docx"
        exp.assemble_docx(root, {}, out)
        from docx import Document
        doc = Document(str(out))
        # 调试：列出所有段落与表格
        all_text = " ".join(p.text for p in doc.paragraphs)
        for t in doc.tables:
            all_text += " [TABLE: " + t.cell(0, 0).text + "]"
        # 风险卡会生成一个表格（含字段）
        self.assertEqual(len(doc.tables), 1, f"实际内容: {all_text!r}")
        cell_text = doc.tables[0].cell(0, 0).text
        self.assertIn("问题描述", cell_text)
        self.assertIn("测试风险", cell_text)

    def test_kpi_card_paragraph(self):
        html = '<div id="sr-page-source"><div class="sr-kpi-card">关键指标 95 分</div></div>'
        doc = self._assemble(html)
        texts = [p.text for p in doc.paragraphs if p.text]
        self.assertIn("关键指标 95 分", texts)

    def test_tag_high_color_run(self):
        html = '<div id="sr-page-source"><span class="sr-tag sr-tag--high">高风险</span></div>'
        doc = self._assemble(html)
        # 至少有一段包含 "高风险"
        texts = " ".join(p.text for p in doc.paragraphs)
        self.assertIn("高风险", texts)

    def test_grade_badge_centered(self):
        html = '<div id="sr-page-source"><div class="sr-grade">良</div></div>'
        doc = self._assemble(html)
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        found = False
        for p in doc.paragraphs:
            if p.text.strip() == "良":
                self.assertEqual(p.alignment, WD_ALIGN_PARAGRAPH.CENTER)
                found = True
        self.assertTrue(found, "未找到评级徽章段落")

    def test_bullet_list(self):
        html = '<div id="sr-page-source"><ul><li>项目一</li><li>项目二</li></ul></div>'
        doc = self._assemble(html)
        bullet_count = sum(1 for p in doc.paragraphs if p.style.name == "List Bullet")
        self.assertEqual(bullet_count, 2)

    def test_number_list(self):
        html = '<div id="sr-page-source"><ol><li>步骤一</li><li>步骤二</li></ol></div>'
        doc = self._assemble(html)
        number_count = sum(1 for p in doc.paragraphs if p.style.name == "List Number")
        self.assertEqual(number_count, 2)

    def test_skip_selectors_excluded(self):
        html = '''<div id="sr-page-source">
            <div class="sr-nav-wrap">NAV_ONLY_TEXT</div>
            <p>CONTENT_TEXT</p>
        </div>'''
        exp = _make_exporter(self.tmpdir)
        exp._snapshot_map = {}
        root = BeautifulSoup(html, "lxml").find(id="sr-page-source")
        out = Path(self.tmpdir) / "skip.docx"
        exp.assemble_docx(root, {}, out)
        from docx import Document
        doc = Document(str(out))
        texts = " ".join(p.text for p in doc.paragraphs)
        self.assertNotIn("NAV_ONLY_TEXT", texts)
        self.assertIn("CONTENT_TEXT", texts)

    def test_table_with_invalid_colspan_falls_back_to_1(self):
        """colspan 非数字时应兜底为 1，不抛异常。"""
        html = '''<div id="sr-page-source"><table>
            <tr><th>A</th><th>B</th></tr>
            <tr><td colspan="abc">合并</td><td>x</td></tr>
        </table></div>'''
        exp = _make_exporter(self.tmpdir)
        exp._snapshot_map = {}
        root = BeautifulSoup(html, "lxml").find(id="sr-page-source")
        out = Path(self.tmpdir) / "bad.docx"
        # 不应抛异常
        exp.assemble_docx(root, {}, out)
        from docx import Document
        doc = Document(str(out))
        self.assertEqual(len(doc.tables), 1)

    def test_snapshot_missing_file_does_not_crash(self):
        """snapshot_map 指向不存在文件时应跳过，不抛异常。"""
        html = '<div id="sr-page-source"><img data-snapshot="bad-snap" alt="x"/></div>'
        exp = _make_exporter(self.tmpdir)
        exp._snapshot_map = {"bad-snap": str(Path(self.tmpdir) / "nonexistent.png")}
        root = BeautifulSoup(html, "lxml").find(id="sr-page-source")
        out = Path(self.tmpdir) / "bad2.docx"
        exp.assemble_docx(root, exp._snapshot_map, out)
        from docx import Document
        doc = Document(str(out))
        # 图片应被跳过，不产生 inline_shape
        self.assertEqual(len(doc.inline_shapes), 0)


class TestSnapshotComponents(_TestBase):
    def test_snapshot_complex_components_uses_selectors(self):
        exp = _make_exporter(self.tmpdir)
        # 只配置 1 个 selector 简化断言
        exp.config["snapshot_selectors"] = [".sr-chart"]
        page = MagicMock()
        locator = MagicMock()
        locator.count.return_value = 2
        el = MagicMock()
        locator.nth.return_value = el
        page.locator.return_value = locator
        page.evaluate.return_value = None
        el.screenshot.return_value = None
        snap_map = exp.snapshot_complex_components(page)
        # 1 selector × 2 元素 = 2 次截图
        self.assertEqual(el.screenshot.call_count, 2)
        self.assertEqual(len(snap_map), 2)

    def test_snapshot_complex_components_handles_screenshot_failure(self):
        exp = _make_exporter(self.tmpdir)
        exp.config["snapshot_selectors"] = [".sr-chart"]
        page = MagicMock()
        locator = MagicMock()
        locator.count.return_value = 1
        el = MagicMock()
        el.screenshot.side_effect = RuntimeError("boom")
        locator.nth.return_value = el
        page.locator.return_value = locator
        page.evaluate.return_value = None
        # 失败时应不抛异常，返回空 map
        snap_map = exp.snapshot_complex_components(page)
        self.assertEqual(snap_map, {})


class TestRunE2EWithMock(_TestBase):
    def test_run_e2e_smoke(self):
        """全流程 mock 烟雾测试，验证 docx 文件生成且非空。"""
        exp = _make_exporter(self.tmpdir)
        # mock 整个 playwright 链
        with patch("html_to_word.html_to_word_export.HtmlToWordExporter.render_html") as mock_render, \
             patch("html_to_word.html_to_word_export.HtmlToWordExporter.snapshot_complex_components") as mock_snap, \
             patch("html_to_word.html_to_word_export.HtmlToWordExporter.extract_dom") as mock_extract, \
             patch("html_to_word.html_to_word_export.HtmlToWordExporter._update_fields_via_word") as mock_upd:
            # render_html 返回一个 mock page，snapshot 返回空 map，extract_dom 返回 fixture DOM
            mock_render.return_value = MagicMock()
            mock_snap.return_value = {}
            html = MINI_HTML.read_text(encoding="utf-8")
            mock_extract.return_value = BeautifulSoup(html, "lxml").find(id="sr-page-source")
            mock_upd.return_value = False
            exp.run()
        # 验证输出文件存在且非空
        self.assertTrue(exp.output_path.exists())
        self.assertGreater(exp.output_path.stat().st_size, 1000)


class TestUpdateFieldsViaWord(_TestBase):
    """验证 _update_fields_via_word 的 COM 编排：配置开关、异常兜底、
    Quit 一定被调用（防 Word 进程泄漏）。"""

    def _make_docx(self):
        """生成一个最小可用的 .docx 文件作为 COM Open 的目标。"""
        from docx import Document
        out = Path(self.tmpdir) / "target.docx"
        Document().save(str(out))
        return out

    def test_disabled_when_config_off(self):
        exp = _make_exporter(self.tmpdir)
        exp.config["update_fields_after_save"] = {
            "enabled": False, "backend": "win32com", "fail_silently": True,
        }
        out = self._make_docx()
        with patch("html_to_word.html_to_word_export.os.open") as mock_open:
            self.assertFalse(exp._update_fields_via_word(out))
            mock_open.assert_not_called()

    def test_skip_when_backend_none(self):
        exp = _make_exporter(self.tmpdir)
        exp.config["update_fields_after_save"] = {
            "enabled": True, "backend": "none", "fail_silently": True,
        }
        out = self._make_docx()
        self.assertFalse(exp._update_fields_via_word(out))

    def test_skip_when_file_locked(self):
        exp = _make_exporter(self.tmpdir)
        exp.config["update_fields_after_save"] = {
            "enabled": True, "backend": "win32com", "fail_silently": True,
        }
        out = self._make_docx()
        with patch("html_to_word.html_to_word_export.os.open",
                   side_effect=OSError("locked")):
            self.assertFalse(exp._update_fields_via_word(out))

    def test_fail_silently_when_no_pywin32(self):
        exp = _make_exporter(self.tmpdir)
        exp.config["update_fields_after_save"] = {
            "enabled": True, "backend": "win32com", "fail_silently": True,
        }
        out = self._make_docx()
        with patch("builtins.__import__",
                   side_effect=ImportError("no pywin32")):
            # 不应抛出，仅返回 False
            self.assertFalse(exp._update_fields_via_word(out))

    def _patch_pywin32(self, fake_w32c, fake_pythoncom=None):
        """统一 patch pywin32 模块链：win32com.client → fake_w32c，
        pythoncom → fake_pythoncom，并保证 win32com.client 属性解析正确。"""
        fake_win32com = MagicMock()
        fake_win32com.client = fake_w32c
        mods = {
            "win32com": fake_win32com,
            "win32com.client": fake_w32c,
        }
        if fake_pythoncom is not None:
            mods["pythoncom"] = fake_pythoncom
        return patch.dict(sys.modules, mods)

    def test_com_pipeline_and_quit_called(self):
        """正常路径：DispatchEx → Open → Fields.Update → TOC.Update →
        Repaginate → Save → Close → Quit 全部执行。"""
        exp = _make_exporter(self.tmpdir)
        exp.config["update_fields_after_save"] = {
            "enabled": True, "backend": "win32com",
            "fail_silently": True, "timeout_sec": 10,
        }
        out = self._make_docx()

        fake_word = MagicMock()
        fake_doc = MagicMock()
        fake_toc = MagicMock()
        fake_tocs = MagicMock(return_value=fake_toc)  # tocs(1) → fake_toc
        fake_tocs.Count = 1
        fake_doc.TablesOfContents = fake_tocs
        fake_word.Documents.Open.return_value = fake_doc

        fake_w32c = MagicMock()
        fake_w32c.DispatchEx.return_value = fake_word
        fake_pythoncom = MagicMock()

        with self._patch_pywin32(fake_w32c, fake_pythoncom):
            result = exp._update_fields_via_word(out)

        self.assertTrue(result)
        fake_w32c.DispatchEx.assert_called_once_with("Word.Application")
        fake_word.Documents.Open.assert_called_once_with(str(out), ReadOnly=False)
        fake_doc.Fields.Update.assert_called_once()
        fake_toc.Update.assert_called_once()
        fake_doc.Repaginate.assert_called_once()
        fake_doc.Save.assert_called_once()
        fake_doc.Close.assert_called_once()
        fake_word.Quit.assert_called_once()
        fake_pythoncom.CoInitialize.assert_called_once()
        fake_pythoncom.CoUninitialize.assert_called_once()

    def test_quit_called_even_on_exception(self):
        """COM Open 抛异常时，Quit 也必须在 finally 中被调用。"""
        exp = _make_exporter(self.tmpdir)
        exp.config["update_fields_after_save"] = {
            "enabled": True, "backend": "win32com", "fail_silently": True,
        }
        out = self._make_docx()

        fake_word = MagicMock()
        fake_word.Documents.Open.side_effect = RuntimeError("open failed")

        fake_w32c = MagicMock()
        fake_w32c.DispatchEx.return_value = fake_word
        fake_pythoncom = MagicMock()

        with self._patch_pywin32(fake_w32c, fake_pythoncom):
            result = exp._update_fields_via_word(out)

        self.assertFalse(result)
        fake_word.Quit.assert_called_once()
        fake_pythoncom.CoUninitialize.assert_called_once()

    def test_raise_when_fail_silently_false(self):
        exp = _make_exporter(self.tmpdir)
        exp.config["update_fields_after_save"] = {
            "enabled": True, "backend": "win32com", "fail_silently": False,
        }
        out = self._make_docx()

        fake_word = MagicMock()
        fake_word.Documents.Open.side_effect = RuntimeError("boom")
        fake_w32c = MagicMock()
        fake_w32c.DispatchEx.return_value = fake_word

        with self._patch_pywin32(fake_w32c, MagicMock()):
            with self.assertRaises(RuntimeError):
                exp._update_fields_via_word(out)
        # 异常路径也要 Quit
        fake_word.Quit.assert_called_once()


if __name__ == "__main__":
    unittest.main(verbosity=2)
