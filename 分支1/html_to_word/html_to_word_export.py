"""
通用 HTML → Word 转换
=====================
将 standalone HTML 报告（含 ECharts / iframe 动态内容）转换为 .docx 文档。

采用「混合保真」策略：
  - 标题 / 段落 / 表格 / 普通图片 → 结构化映射到 Word 原生元素（可编辑）
  - ECharts 图表 / iframe sketch 图 / JS 动态 SVG → 截图嵌入（保视觉）

用法:
    # 默认转换（A4 横版分页）
    python -m html_to_word.html_to_word_export --input "security-report-preview-2.0-standalone(3).html"

    # 指定输出路径
    python -m html_to_word.html_to_word_export --input xxx.html --output xxx.docx

    # 指定预览模式（默认 a4-landscape）
    python -m html_to_word.html_to_word_export --input xxx.html --preview-mode a4-portrait

    # 自定义配置文件
    python -m html_to_word.html_to_word_export --input xxx.html --config html_to_word/html_to_word_config.yaml
"""
import time
import argparse
import base64
import io
import os
import re
from pathlib import Path

import yaml
from bs4 import BeautifulSoup, Comment, NavigableString, Tag
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor, Emu
from PIL import Image

# 风险卡头部标签与卡片上缘的间距（Word 单元格无内边距，用 space_before 模拟 HTML 的 padding-top）
# 标签本体保持可编辑文本（run 底纹），仅解决"顶格、无上间距"问题。
_RISK_CARD_HEAD_TAG_SPACE_BEFORE = 8.0   # pt

# 默认配置（与 html_to_word_config.yaml 同步，配置缺失时回退到此）
_DEFAULT_CONFIG = {
    "preview_mode": "html",
    "trigger_paginate": False,
    "render_wait": {"goto_timeout_ms": 60000, "after_load_ms": 2000, "ready_marker_timeout_ms": 30000, "ready_markers": []},
    "dom_root": "#sr-page-source",
    "skip_selectors": [".sr-nav-wrap", ".sr-toc-page", ".sr-copyright-page", ".sr-dev-reference", ".sr-hidden-charts"],
    "snapshot_selectors": [
        ".sr-chart-slot",
        ".sr-chart-card",
        ".sr-chart-live",
        ".sr-asset-donut-row",
        "iframe",
        ".sr-chart-empty",
    ],
    "style_map": {
        "h1": "heading_1", "h2": "heading_2", "h3": "heading_3",
        "h4": "heading_4", "h5": "heading_5", "h6": "heading_6",
        "p": "paragraph", "table": "table",
        "ul": "bullet_list", "ol": "number_list", "img": "image",
        ".sr-risk-card": "risk_card_table",
        ".sr-risk-card__head": "risk_card_head",
        ".sr-kpi-card": "kpi_paragraph",
        ".sr-top5-asset-ip-row": "top5_asset_ip_row",
        ".sr-tag--critical": "tag_critical",
        ".sr-tag--high": "tag_high",
        ".sr-tag--medium": "tag_medium",
        ".sr-tag--medium-low": "tag_medium_low",
        ".sr-tag--low": "tag_low",
        ".sr-tag--info": "tag_info",
        ".sr-tag--success": "tag_success",
        ".sr-tag--blue": "tag_blue",
        ".sr-tag--light": "tag_light",
        ".sr-grade": "grade_badge",
        ".report-chart-note": "muted_note",
        ".sr-p--sub": "sub_paragraph",
        ".sr-risk-harm-box": "risk_harm_box",
    },
    "docx": {"page_width_mm": 210, "page_height_mm": 297, "margin_mm": 20},
    "fonts": {
        "normal": {"font": "Microsoft YaHei", "east_asia": "Microsoft YaHei",
                   "size_pt": 10.5, "color_hex": "000000",
                   "space_after_pt": 6},
        "heading_1": {"font": "Microsoft YaHei", "east_asia": "Microsoft YaHei",
                      "size_pt": 15, "bold": True, "color_hex": "1F4E79",
                      "space_before_pt": 18, "space_after_pt": 9},
        "heading_2": {"font": "Microsoft YaHei", "east_asia": "Microsoft YaHei",
                      "size_pt": 14, "bold": True, "color_hex": "1F4E79",
                      "space_before_pt": 18, "space_after_pt": 9},
        "heading_3": {"font": "Microsoft YaHei", "east_asia": "Microsoft YaHei",
                      "size_pt": 10.5, "bold": True, "color_hex": "2E5C8A",
                      "space_before_pt": 6, "space_after_pt": 9},
        "heading_4": {"font": "Microsoft YaHei", "east_asia": "Microsoft YaHei",
                      "size_pt": 10.5, "bold": True, "color_hex": "2E5C8A",
                      "space_before_pt": 6, "space_after_pt": 9},
        "heading_5": {"font": "Microsoft YaHei", "east_asia": "Microsoft YaHei",
                      "size_pt": 10.5, "bold": True, "color_hex": "44546A",
                      "space_before_pt": 6, "space_after_pt": 9},
        "heading_6": {"font": "Microsoft YaHei", "east_asia": "Microsoft YaHei",
                      "size_pt": 10.5, "bold": True, "color_hex": "44546A",
                      "space_before_pt": 6, "space_after_pt": 9},
    },
    "screenshot": {"type": "png", "scale": 2},
    "tmp_dir": "tmp/html_to_word",
    "update_fields_after_save": {
        "enabled": True,
        "backend": "win32com",
        "timeout_sec": 60,
        "fail_silently": True,
    },
}

# 风险标签色阶（CSS 变量对应色，Word 中用近似 RGB）
_TAG_COLORS = {
    "tag_critical": ("82010E", "FFFFFF"),
    "tag_high":     ("CF171D", "FFFFFF"),
    "tag_medium":   ("FA721B", "FFFFFF"),
    "tag_medium_low": ("D6860D", "FFFFFF"),
    "tag_low":      ("6F7785", "FFFFFF"),
    "tag_info":     ("0BA7B5", "FFFFFF"),
    "tag_success":  ("12A679", "FFFFFF"),
    "tag_blue":     ("1C6EFF", "FFFFFF"),
}

# 评级徽章 sr-grade 颜色映射
# HTML 中 sr-grade--xxx 用 10% 透明底 + 深色字，Word w:shd 无透明度，用浅 solid 底近似
# (背景色, 字体颜色)
_GRADE_COLORS = {
    "sr-grade--优":  ("E5F5EE", "0F8E66"),   # 绿色徽章
    "sr-grade--良":  ("E6F0FF", "1C6EFF"),   # 蓝色徽章
    "sr-grade--中":  ("FEEAD8", "B5530C"),   # 橙色徽章
    "sr-grade--差":  ("FBE5E6", "82010E"),   # 红色徽章
    "sr-grade--na":  ("EDEEF1", "6F7785"),   # 灰色徽章
}
_GRADE_DEFAULT = ("EDEEF1", "6F7785")


def _log(msg, level="INFO"):
    print(f"[{level}] {msg}")


def _has_class(node, class_name):
    """判断 BeautifulSoup 节点是否含指定 class。

    Args:
        node: bs4 Tag / NavigableString 等。非 Tag 一律返回 False。
        class_name: 待匹配的 class 名（精确匹配，多个 class 时任一命中即 True）。
    """
    if not isinstance(node, Tag):
        return False
    classes = node.get("class") or []
    if isinstance(classes, str):
        classes = classes.split()
    return class_name in classes


class _Container:
    """统一封装 Document / Cell 的添加接口。

    - Document：直接用 add_heading / add_paragraph / add_table / add_picture
    - _Cell：用 add_paragraph / add_table；图片与标题用变体实现
    """

    def __init__(self, doc, cell=None, fonts=None):
        self.doc = doc
        self.cell = cell  # 非 None 表示容器是 Cell
        # 字体样式配置（来自 HtmlToWordExporter.config["fonts"]），
        # 用于单元格内标题字号/颜色/东亚字体回退。
        self.fonts = fonts or {}

    @property
    def is_cell(self):
        return self.cell is not None

    def add_heading(self, text, level):
        if self.is_cell:
            p = self.cell.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            # 优先读 fonts.heading_<level> 配置；缺省回退到内置字号映射
            heading_cfg = (self.fonts.get(f"heading_{level}", {}) or {}) if self.fonts else {}
            default_size_map = {1: 15, 2: 14, 3: 10.5, 4: 10.5, 5: 10.5, 6: 10.5}
            size_pt = heading_cfg.get("size_pt") or default_size_map.get(level, 11)
            run.font.size = Pt(size_pt)
            if heading_cfg.get("color_hex"):
                try:
                    run.font.color.rgb = RGBColor.from_string(heading_cfg["color_hex"])
                except (ValueError, AttributeError):
                    pass
            # 东亚字体：heading 优先，回退到 normal.east_asia
            ea = heading_cfg.get("east_asia") or (self.fonts.get("normal", {}) or {}).get("east_asia")
            latin = heading_cfg.get("font") or (self.fonts.get("normal", {}) or {}).get("font")
            if ea or latin:
                rPr = run._r.get_or_add_rPr()
                rFonts = rPr.find(qn("w:rFonts"))
                if rFonts is None:
                    rFonts = OxmlElement("w:rFonts")
                    rPr.append(rFonts)
                # 清除主题字体引用，避免 Word fallback 到主题字体（MS Gothic 等）
                for theme_attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
                    if rFonts.get(qn(f"w:{theme_attr}")) is not None:
                        del rFonts.attrib[qn(f"w:{theme_attr}")]
                if ea:
                    rFonts.set(qn("w:eastAsia"), ea)
                if latin:
                    rFonts.set(qn("w:ascii"), latin)
                    rFonts.set(qn("w:hAnsi"), latin)
            return p
        return self.doc.add_heading(text, level=level)

    def add_paragraph(self, text="", style=None):
        if self.is_cell:
            return self.cell.add_paragraph(text, style=style)
        return self.doc.add_paragraph(text, style=style)

    def add_table(self, rows, cols):
        if self.is_cell:
            return self.cell.add_table(rows, cols)
        return self.doc.add_table(rows, cols)

    def add_picture(self, image, width=None, height=None):
        if self.is_cell:
            # Cell 没有原生 add_picture：在新建段落里用 run.add_picture
            p = self.cell.add_paragraph()
            run = p.add_run()
            if isinstance(image, str):
                run.add_picture(image, width=width, height=height)
            else:
                run.add_picture(image, width=width, height=height)
            return p
        # Document：python-docx 的 add_picture 返回 InlineShape，无法直接拿段落；
        # 改为手动建段，让调用方能拿到段落对象设置 space_before / space_after。
        p = self.doc.add_paragraph()
        run = p.add_run()
        if isinstance(image, str):
            run.add_picture(image, width=width, height=height)
        else:
            run.add_picture(image, width=width, height=height)
        return p


def _css_to_bs4_selector(selector):
    """简单将 .cls / #id / tag 选择器转换为 bs4 find 参数。"""
    selector = selector.strip()
    if selector.startswith("#"):
        return {"id": selector[1:]}
    if selector.startswith("."):
        return {"class_": selector[1:].replace(".", " ")}
    return {"name": selector}


def _match_any(tag, selectors):
    """判断 tag 是否匹配任一选择器（仅支持 .cls / #id / tag 简单形式）。"""
    if not isinstance(tag, Tag):
        return False
    for sel in selectors:
        sel = sel.strip()
        if sel.startswith("#"):
            if tag.get("id") == sel[1:]:
                return True
        elif sel.startswith("."):
            classes = (tag.get("class") or [])
            needed = sel[1:].split(".")
            if all(c in classes for c in needed):
                return True
        else:
            if tag.name == sel:
                return True
    return False


class HtmlToWordExporter:
    """HTML → Word 转换器：渲染、截图、结构化映射、拼装 docx。"""

    def __init__(self, input_path, output_path=None, config_path=None, preview_mode=None):
        self.input_path = Path(input_path).resolve()
        if not self.input_path.exists():
            raise FileNotFoundError(f"输入 HTML 不存在: {self.input_path}")
        self.output_path = Path(output_path) if output_path else (
            Path("tmp") / (self.input_path.stem + ".docx")
        )
        # 默认使用脚本同目录下的 html_to_word_config.yaml，避免 main() 不传 --config 时
        # 直接落到内置 _DEFAULT_CONFIG，导致 yaml 中的 ready_markers 等修正不生效。
        default_cfg = Path(__file__).resolve().parent / "html_to_word_config.yaml"
        self.config_path = Path(config_path) if config_path else (
            default_cfg if default_cfg.exists() else None
        )
        self.config = dict(_DEFAULT_CONFIG)
        self._snapshot_index = 0
        self._snapshot_map = {}  # DOM 占位 id → 图片绝对路径
        # 延迟导入 playwright，避免无浏览器环境也能跑测试
        self._playwright_ctx = None
        self._page = None
        # 封面节点缓存（.sr-cover 在 dom_root 之外，extract_dom 时填充）
        self._cover_node = None
        # preview_mode 覆盖
        if preview_mode:
            self.config["preview_mode"] = preview_mode

    # ──────────────────────────────────────────────
    # 配置加载
    # ──────────────────────────────────────────────

    def load_config(self):
        """加载 YAML 配置，与默认配置深合并。"""
        if not self.config_path or not self.config_path.exists():
            _log(f"未指定配置文件，使用内置默认配置")
            return self.config
        with open(self.config_path, "r", encoding="utf-8") as f:
            user_cfg = yaml.safe_load(f) or {}
        # 深合并：用户配置覆盖默认配置
        for k, v in user_cfg.items():
            if isinstance(v, dict) and isinstance(self.config.get(k), dict):
                self.config[k] = {**self.config[k], **v}
            else:
                self.config[k] = v
        _log(f"已加载配置: {self.config_path}")
        return self.config

    # ──────────────────────────────────────────────
    # 渲染（Playwright）
    # ──────────────────────────────────────────────

    def render_html(self):
        """Playwright 启动 Chromium，加载 HTML，设置 preview-mode，等待渲染完成。返回 page。"""
        from playwright.sync_api import sync_playwright
        _log(f"启动 Chromium 渲染: {self.input_path.name}")
        self._playwright_ctx = sync_playwright().start()
        # 优先用系统 Chrome（避免依赖 playwright 自带 chromium 下载）
        chrome_path = self._find_system_chrome()
        launch_kwargs = {}
        if chrome_path:
            launch_kwargs["executable_path"] = chrome_path
            _log(f"使用系统 Chrome: {chrome_path}")
        try:
            browser = self._playwright_ctx.chromium.launch(**launch_kwargs)
        except Exception as e:
            _log(f"系统 Chrome 启动失败，回退到自带 chromium: {e}", "WARNING")
            browser = self._playwright_ctx.chromium.launch()
        context = browser.new_context(
            viewport={"width": 1600, "height": 900},
            device_scale_factor=self.config.get("screenshot", {}).get("scale", 2),
        )
        page = context.new_page()
        url = self.input_path.as_uri()
        wait_cfg = self.config.get("render_wait", {})
        goto_timeout_ms = wait_cfg.get("goto_timeout_ms", 60000)
        page.goto(url, wait_until="load", timeout=goto_timeout_ms)
        # 注入 preview-mode
        preview_mode = self.config.get("preview_mode", "a4-landscape")
        page.evaluate(
            f"localStorage.setItem('sr-preview-mode', '{preview_mode}')"
        )
        page.reload(wait_until="load", timeout=goto_timeout_ms)
        # 触发分页（可选）
        if self.config.get("trigger_paginate", False):
            page.evaluate("window.paginate && window.paginate()")
        # 等待渲染完成
        after_load_ms = wait_cfg.get("after_load_ms", 2000)
        marker_timeout_ms = wait_cfg.get("ready_marker_timeout_ms", 30000)
        page.wait_for_timeout(after_load_ms)
        for marker in wait_cfg.get("ready_markers", []):
            try:
                page.wait_for_selector(marker, timeout=marker_timeout_ms)
            except Exception as e:
                _log(f"ready_marker 未出现（忽略）: {marker} ({e})", "WARNING")
        self._page = page
        self._browser = browser
        self._context = context
        _log("渲染完成")
        return page

    @staticmethod
    def _find_system_chrome():
        """查找系统已装的 Chrome / Edge 可执行文件路径。"""
        candidates = [
            r"C:\Program Files\Google\Chrome\Application\chrome.exe",
            r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
            r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
            r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
        ]
        for p in candidates:
            if os.path.exists(p):
                return p
        return None

    def close(self):
        if self._playwright_ctx:
            try:
                self._browser.close()
            except Exception:
                pass
            try:
                self._playwright_ctx.stop()
            except Exception:
                pass
            self._playwright_ctx = None

    # ──────────────────────────────────────────────
    # 复杂组件截图
    # ──────────────────────────────────────────────

    @staticmethod
    def _freeze_layout_script():
        """注入脚本：收尾 ECharts 渲染 + 把 .chw 容器尺寸固定为内联样式 + 冻结
        ResizeObserver + 隐藏 nav，避免 Playwright `Locator.screenshot` 在
        scrollIntoView 后元素持续重排导致 30s 超时。

        根因（实测确认）：已存在的 ResizeObserver 实例引用的是原构造器，替换
        `window.ResizeObserver` 不能阻止它们触发 callback；唯一可靠办法是让被
        观察元素的尺寸不再变化。本脚本读出每个 `.chw` 当前的 offsetWidth/Height，
        写为内联 `width/height`，使后续 RO 测得的尺寸恒等于当前值，回调不会再
        触发 echarts 重排。配合隐藏 nav，scrollIntoView 也不会触发 onScroll 修改
        DOM。
        """
        return r"""
        (() => {
          const ret = { resized: 0, observersDisconnected: 0, roReplaced: false,
                        navHidden: false, chwPinned: 0 };
          try {
            // 1) 主动收尾：把所有 .chw 容器上的 ECharts 实例 resize 一次，
            //    让 size 落到当前布局的最终值。
            if (window.echarts) {
              document.querySelectorAll('.chw').forEach((el) => {
                try {
                  const inst = window.echarts.getInstanceByDom(el);
                  if (inst) { inst.resize(); ret.resized += 1; }
                } catch (e) { /* ignore individual chart error */ }
              });
            }
            // 2) 断开已知 donut chart 的 ResizeObserver（__srDonutCleanup）。
            if (window.echarts) {
              document.querySelectorAll('.chw').forEach((el) => {
                try {
                  const inst = window.echarts.getInstanceByDom(el);
                  if (inst && typeof inst.__srDonutCleanup === 'function') {
                    inst.__srDonutCleanup();
                    ret.observersDisconnected += 1;
                  }
                } catch (e) { /* ignore */ }
              });
            }
            // 3) 冻结后续 ResizeObserver：替换 window.ResizeObserver 为 no-op。
            if (window.ResizeObserver) {
              window.__srOrigResizeObserver = window.ResizeObserver;
              window.ResizeObserver = class {
                constructor(cb) { this._cb = cb; }
                observe() {}
                unobserve() {}
                disconnect() {}
              };
              ret.roReplaced = true;
            }
            // 4) 隐藏 nav（nav 本来就在 skip_selectors 中），setActive 即使被
            //    onScroll 调用也无 DOM 可改 class。
            try {
              document.querySelectorAll('.sr-nav-wrap, .sr-nav-rail, .sr-nav').forEach((el) => {
                el.style.display = 'none';
              });
              ret.navHidden = true;
            } catch (e) { /* ignore */ }
            // 5) 关键：把每个 .chw 容器尺寸固定为内联样式。这样被观察元素的
            //    尺寸不再变化，已存在的 RO 实例不会再触发 callback（即使它们
            //    引用的是原构造器）。再调用一次 resize() 让 ECharts 内部画布
            //    与固定后的容器尺寸对齐。
            document.querySelectorAll('.chw').forEach((el) => {
              try {
                const w = el.offsetWidth;
                const h = el.offsetHeight;
                if (w > 0 && h > 0) {
                  el.style.width = w + 'px';
                  el.style.height = h + 'px';
                  ret.chwPinned += 1;
                }
              } catch (e) { /* ignore */ }
            });
            if (window.echarts) {
              document.querySelectorAll('.chw').forEach((el) => {
                try {
                  const inst = window.echarts.getInstanceByDom(el);
                  if (inst) inst.resize();
                } catch (e) { /* ignore */ }
              });
            }
            // 6) 滚回顶部，避免 scrollIntoView 大幅滚动触发更多 listener。
            try { window.scrollTo(0, 0); } catch (e) { /* ignore */ }
          } catch (e) {
            ret.error = String(e);
          }
          return ret;
        })();
        """

    def _freeze_layout_before_snapshot(self, page):
        """截图前冻结布局：收尾 ECharts + 固定 .chw 尺寸 + 隐藏 nav。"""
        try:
            result = page.evaluate(self._freeze_layout_script())
            _log(f"冻结布局: resized={result.get('resized')} "
                 f"observersDisconnected={result.get('observersDisconnected')} "
                 f"roReplaced={result.get('roReplaced')} "
                 f"navHidden={result.get('navHidden')} "
                 f"chwPinned={result.get('chwPinned')} "
                 f"error={result.get('error', 'none')}")
            # 让最后一帧 layout 跑完，再开始截图
            page.wait_for_timeout(300)
        except Exception as e:
            _log(f"冻结布局注入失败（忽略）: {e}", "WARNING")

    def _snapshot_grid2_in_slot(self, page, slot_snap_id, tmp_dir, snapshot_map):
        """把 slot 内部第一个 .sr-chart-grid-2 整体截图成一张图，
        在 DOM 中替换为 <img data-snapshot> 占位，写入 snapshot_map。
        用于 3.3.1 互联网业务 slot：把 Web top5 + 非Web top5 合成一张图。"""
        grid_snap_id = f"sr-snap-{self._snapshot_index}"
        self._snapshot_index += 1
        # 给 slot 内的 grid-2 打标
        marked = page.evaluate(
            """([slotSid, gridSid]) => {
                const slot = document.querySelector(`[data-snapshot-id="${slotSid}"]`);
                if (!slot) return false;
                const grid = slot.querySelector('.sr-chart-grid-2');
                if (!grid) return false;
                grid.setAttribute('data-snapshot-id', gridSid);
                return true;
            }""",
            [slot_snap_id, grid_snap_id],
        )
        if not marked:
            return
        grid_locator = page.locator(f'[data-snapshot-id="{grid_snap_id}"]')
        img_path = tmp_dir / f"comp-{grid_snap_id}.png"
        grid_locator.screenshot(path=str(img_path))
        # 替换为 img 占位
        page.evaluate(
            """(sid) => {
                const el = document.querySelector(`[data-snapshot-id="${sid}"]`);
                if (!el) return;
                const img = document.createElement('img');
                img.setAttribute('data-snapshot', sid);
                img.setAttribute('alt', 'chart');
                img.style.maxWidth = '100%';
                el.replaceWith(img);
            }""",
            grid_snap_id,
        )
        snapshot_map[grid_snap_id] = str(img_path)
        # 记录 grid-2 合成图的图注文字（图下居中显示）
        if not hasattr(self, "_snapshot_grid_captions"):
            self._snapshot_grid_captions = {}
        self._snapshot_grid_captions[grid_snap_id] = "Web服务和非Web服务风险分布 top5"

    def snapshot_complex_components(self, page=None):
        """遍历 snapshot_selectors，逐个截图，返回 {占位 id: 图片路径} 映射。"""
        page = page or self._page
        # 截图前先冻结一次布局：把 ECharts 收尾 + 固定 .chw 容器尺寸 + 隐藏 nav。
        # 循环内每次截图前还会再冻结一次，以应对上一轮 replaceWith 引起的 DOM 变化。
        self._freeze_layout_before_snapshot(page)
        tmp_dir = Path(self.config.get("tmp_dir", "tmp/html_to_word"))
        tmp_dir.mkdir(parents=True, exist_ok=True)
        snapshot_map = {}
        for selector in self.config.get("snapshot_selectors", []):
            try:
                locator = page.locator(selector)
                initial_count = locator.count()
            except Exception as e:
                _log(f"selector 不可用（忽略）: {selector} ({e})", "WARNING")
                continue
            _log(f"截图 selector={selector} count={initial_count}")
            # 关键：每次截图后用 replaceWith 把元素替换为 img 占位，会让该 selector
            # 对应的元素列表少一个。若仍按 range(initial_count) 迭代，第 i 越往后
            # 越容易取不到元素（els[idx] 为 undefined）。改用 while 循环：每次取
            # 第 0 个元素、打标、截图、替换，直到该 selector 下无元素为止。
            # 注意：嵌套选择器（如 .sr-chart-card 嵌在 .sr-chart-slot 内）会随外层
            # 替换而消失，循环会自然提前结束——这是预期行为（外层 slot 已截图覆盖）。
            captured = 0
            while True:
                # 截图前先冻结布局：上一轮 replaceWith 会改变 DOM 结构，可能导致
                # 剩余 .chw 容器尺寸变化、ECharts 重排。重新跑一次冻结脚本把剩余
                # 容器尺寸固定到当前值，确保截图时 bounding box 稳定。
                self._freeze_layout_before_snapshot(page)
                snap_id = f"sr-snap-{self._snapshot_index}"
                self._snapshot_index += 1
                # 给当前 selector 下第 0 个元素打 data-snapshot-id 标记
                try:
                    marked = page.evaluate(
                        """([sel, sid]) => {
                            const el = document.querySelector(sel);
                            if (!el) return false;
                            el.setAttribute('data-snapshot-id', sid);
                            return true;
                        }""",
                        [selector, snap_id],
                    )
                except Exception as e:
                    _log(f"打标失败（忽略）: {selector} ({e})", "WARNING")
                    break
                if not marked:
                    # 该 selector 下已无元素，结束循环
                    break

                # .sr-chart-slot 仅含 .sr-chart-card / 布局网格时跳过整槽截图，
                # 让内部卡片单独截图（Word 里多张分布图独立呈现）
                # 例外 1：id=slot-event-charts（3.2.3 安全事件分布）整体截图成一张图，
                #         保留 HTML 里两张 card 并排 + 各自带标题的视觉，不在 Word 里加图前/图后文字。
                # 例外 2：id=slot-internet-exposure（3.3.1 互联网业务）先把内部 .sr-chart-grid-2
                #         （Web top5 + 非Web top5）整体截图合成一张图替换为 img 占位，
                #         再让 slot 走原跳过逻辑（其它两个 --full card 各自独立截图）。
                if selector == ".sr-chart-slot":
                    try:
                        slot_id = page.evaluate(
                            """(sid) => {
                                const el = document.querySelector(`[data-snapshot-id="${sid}"]`);
                                return el ? (el.getAttribute('id') || '') : '';
                            }""",
                            snap_id,
                        )
                        # 例外 1：整体截图
                        if slot_id == 'slot-event-charts':
                            is_event_charts_slot = True
                        else:
                            is_event_charts_slot = False
                        # 例外 2：先把 grid-2 合成一张图
                        if slot_id == 'slot-internet-exposure':
                            try:
                                self._snapshot_grid2_in_slot(page, snap_id, tmp_dir, snapshot_map)
                            except Exception as e:
                                _log(f"slot-internet-exposure grid-2 合成截图失败（忽略）: {e}", "WARNING")
                        if is_event_charts_slot:
                            # 不拆卡，直接整体截图（落到下面的截图主流程）
                            should_skip = False
                        else:
                            should_skip = page.evaluate(
                                """(sid) => {
                                    const el = document.querySelector(`[data-snapshot-id="${sid}"]`);
                                    if (!el) return false;
                                    const children = Array.from(el.children).filter(c => c.nodeType === 1);
                                    if (children.length === 0) return false;
                                    // 允许直接子元素是 .sr-chart-card / 布局网格容器 / 已截图的 img 占位
                                    // （.sr-chart-grid-2 / .sr-chart-grid-3 等只承载布局；
                                    //  img[data-snapshot] 是 grid-2 被合成截图后替换的占位）
                                    const layoutOnly = c =>
                                        c.classList.contains('sr-chart-card') ||
                                        /sr-chart-grid-\\d/.test(c.className || '') ||
                                        (c.tagName === 'IMG' && c.hasAttribute('data-snapshot'));
                                    const allAllowed = children.every(layoutOnly);
                                    if (!allAllowed) return false;
                                    // 进一步校验：子元素里至少要有一张 .sr-chart-card
                                    // （布局网格内的卡片也算；已截图的 img 占位不算）
                                    const hasCard = children.some(c =>
                                        c.classList.contains('sr-chart-card') ||
                                        c.querySelector('.sr-chart-card'));
                                    if (!hasCard) return false;
                                    el.removeAttribute('data-snapshot-id');
                                    const parent = el.parentNode;
                                    while (el.firstChild) {
                                        parent.insertBefore(el.firstChild, el);
                                    }
                                    parent.removeChild(el);
                                    return true;
                                }""",
                                snap_id,
                            )
                    except Exception as e:
                        _log(f".sr-chart-slot 检查失败（忽略）: {e}", "WARNING")
                        should_skip = False
                    if should_skip:
                        continue
                # 用 data-snapshot-id 精确定位（不受 replaceWith 索引偏移影响）
                el = page.locator(f'[data-snapshot-id="{snap_id}"]')
                img_path = tmp_dir / f"comp-{snap_id}.png"

                # .sr-chart-card 截图前：提取标题（chart-title）。
                # 默认移除标题（让 Word 把标题独立放在图上）；
                # 但 keep_title_in_image 集合中的 chart_id 保留标题在图里（仅写图后居中 caption）。
                # no_caption_ids 集合中的 chart_id 既不写图前小标题、也不写图后居中 caption。
                chart_title = None
                chart_id = None
                keep_title_ids = {"m3-exposure-overview-bar", "m3-bar"}
                no_caption_ids = {"top5-risk-bar"}  # 2.2 风险资产 TOP5 图：标题留图里，无图前/图后文字
                if selector == ".sr-chart-card":
                    try:
                        result = page.evaluate(
                            """([sid, keepIds, noCapIds]) => {
                                const el = document.querySelector(`[data-snapshot-id="${sid}"]`);
                                if (!el) return null;
                                const title = el.querySelector('.chart-title, .sr-chart-title');
                                let titleText = null;
                                let keepTitle = false;
                                // 找内部图表容器 ID
                                const chartEl = el.querySelector('[id]:not([id=""])');
                                const chartId = chartEl ? chartEl.getAttribute('id') : null;
                                if (title) {
                                    titleText = (title.textContent || '').trim();
                                    // chartId 在 keepIds/noCapIds 中时保留标题在图里
                                    if (chartId && (keepIds.includes(chartId) || noCapIds.includes(chartId))) {
                                        keepTitle = true;
                                    } else {
                                        title.parentNode.removeChild(title);
                                    }
                                }
                                return { title: titleText, chartId: chartId, keepTitle: keepTitle };
                            }""",
                            [snap_id, list(keep_title_ids), list(no_caption_ids)],
                        )
                        if result:
                            chart_title = result.get('title')
                            chart_id = result.get('chartId')
                            if result.get('keepTitle') and chart_id:
                                if not hasattr(self, "_snapshot_keep_title"):
                                    self._snapshot_keep_title = set()
                                self._snapshot_keep_title.add(snap_id)
                                if chart_id in no_caption_ids:
                                    if not hasattr(self, "_snapshot_no_caption"):
                                        self._snapshot_no_caption = set()
                                    self._snapshot_no_caption.add(snap_id)
                    except Exception as e:
                        _log(f".sr-chart-card 标题提取失败（忽略）: {e}", "WARNING")
                if chart_title:
                    if not hasattr(self, "_snapshot_titles"):
                        self._snapshot_titles = {}
                    self._snapshot_titles[snap_id] = chart_title
                if chart_id:
                    if not hasattr(self, "_snapshot_chart_ids"):
                        self._snapshot_chart_ids = {}
                    self._snapshot_chart_ids[snap_id] = chart_id

                try:
                    el.screenshot(path=str(img_path))
                except Exception as e:
                    _log(f"截图失败（忽略）: {selector}#{captured} ({e})", "WARNING")
                    # 截图失败也要清理标记，避免影响后续
                    try:
                        page.evaluate(
                            """(sid) => {
                                const el = document.querySelector(`[data-snapshot-id="${sid}"]`);
                                if (el) el.removeAttribute('data-snapshot-id');
                            }""",
                            snap_id,
                        )
                    except Exception:
                        pass
                    # 即使截图失败也把元素替换为空 img，避免下一轮 while 又取到同一元素死循环
                    try:
                        page.evaluate(
                            """([sel, sid]) => {
                                const el = document.querySelector(`[data-snapshot-id="${sid}"]`)
                                    || document.querySelector(sel);
                                if (!el) return;
                                const img = document.createElement('img');
                                img.setAttribute('data-snapshot', sid);
                                img.setAttribute('alt', 'chart');
                                img.style.maxWidth = '100%';
                                el.replaceWith(img);
                            }""",
                            [selector, snap_id],
                        )
                    except Exception:
                        pass
                    continue
                # 替换为 img 占位：通过 data-snapshot-id 定位
                try:
                    page.evaluate(
                        """(sid) => {
                            const el = document.querySelector(`[data-snapshot-id="${sid}"]`);
                            if (!el) return;
                            const img = document.createElement('img');
                            img.setAttribute('data-snapshot', sid);
                            img.setAttribute('alt', 'chart');
                            img.style.maxWidth = '100%';
                            el.replaceWith(img);
                        }""",
                        snap_id,
                    )
                except Exception as e:
                    _log(f"DOM 替换失败（忽略）: {selector}#{captured} ({e})", "WARNING")
                    continue
                snapshot_map[snap_id] = str(img_path)
                captured += 1
            _log(f"selector={selector} 实际截图 {captured}/{initial_count} 个")
        _log(f"截图完成: {len(snapshot_map)} 个组件")
        self._snapshot_map = snapshot_map
        return snapshot_map

    # ──────────────────────────────────────────────
    # DOM 提取
    # ──────────────────────────────────────────────

    def extract_dom(self, page=None):
        """page.content() → BeautifulSoup。返回根节点（按 dom_root 选择）。"""
        page = page or self._page
        html = page.content()
        soup = BeautifulSoup(html, "lxml")
        root_sel = self.config.get("dom_root", "#sr-page-source")
        if root_sel.startswith("#"):
            root = soup.find(id=root_sel[1:])
        else:
            root = soup.select_one(root_sel)
        if root is None:
            _log(f"dom_root 未找到，回退到 <body>", "WARNING")
            root = soup.body or soup
        # 封面 .sr-cover 在 dom_root 之外（HTML 中作为页面顶层 header 存在），
        # 此处先从 soup 抓取并缓存，避免 _extract_cover_info 在 root 内找不到。
        self._cover_node = soup.select_one(".sr-cover")
        # 文档信息页 .sr-doc-info-page 在 dom_root 之内，但要单独成页并排在 TOC 前，
        # 此处先从 soup 抓取并缓存，避免正文流里重复映射。
        self._doc_info_node = soup.select_one(".sr-doc-info-page")
        # 剔除 skip_selectors 节点
        for sel in self.config.get("skip_selectors", []):
            for node in root.select(sel):
                node.decompose()
        # 剔除 HTML 注释（浏览器不渲染，但 NavigableString 分支会把它当文本
        # 提取并写入 Word，导致 "TOP5 风险详情行：#01/#02 typePhrase 动态规则
        # 见 安全体检报告2.0-需求与交互设计评审.md §P4 §4.1 延伸盘点" 这类
        # 设计文档注释出现在正文里）
        for c in root.find_all(string=lambda s: isinstance(s, Comment)):
            c.extract()
        # 剔除 script / style / noscript
        for tag in root.find_all(["script", "style", "noscript"]):
            tag.decompose()
        # 剔除 data-hide="true" 元素（HTML 预览中被 CSS [data-hide="true"]{display:none}
        # 隐藏的内容，如占位符列表项、无需展示的风险卡等，浏览器不渲染，Word 也不应出现）
        for tag in root.find_all(attrs={"data-hide": "true"}):
            tag.decompose()
        return root

    # ──────────────────────────────────────────────
    # docx 拼装
    # ──────────────────────────────────────────────

    def assemble_docx(self, root, snapshot_map, output_path=None):
        """创建 docx，遍历 DOM 映射，保存。"""
        # 防御性剔除 skip_selectors（即使调用方未走 extract_dom）
        for sel in self.config.get("skip_selectors", []):
            try:
                for node in root.select(sel):
                    node.decompose()
            except Exception:
                pass
        for tag in root.find_all(["script", "style", "noscript"]):
            tag.decompose()
        # 防御性剔除 HTML 注释（同 extract_dom）
        for c in root.find_all(string=lambda s: isinstance(s, Comment)):
            c.extract()
        # 防御性剔除 data-hide="true" 元素（同 extract_dom）
        try:
            for tag in root.find_all(attrs={"data-hide": "true"}):
                tag.decompose()
        except Exception:
            pass
        doc = Document()
        self._configure_page(doc)
        # 首页：A4 竖版封面（背景图固定使用 html_to_word/assets/a4-portrait-bg.png）
        # 必须在正文 section 配置之后、TOC 之前插入；插入完成后用 section break
        # 切换回正文 section（A4 横版）。
        cover_info = self._extract_cover_info(root)
        self._insert_cover_page(doc, cover_info)
        # 切换到正文 section：在封面后新增一个 next-page section break，恢复横版
        self._switch_to_body_section(doc)
        # 文档信息页：单独成页，排在 TOC 前
        self._insert_doc_info_page(doc)
        # 在正文 section 开头插入目录（TOC 域）
        self._insert_toc(doc, levels="1-3")
        container = _Container(doc, fonts=self.config.get("fonts", {}) or {})
        self._map_children(root, container)
        out = Path(output_path) if output_path else self.output_path
        out.parent.mkdir(parents=True, exist_ok=True)
        # 在 settings.xml 写入 <w:updateFields w:val="true"/>，Word 打开文档时
        # 会自动提示更新所有域（包含 TOC），用户点"是"即可填充真实目录。
        self._enable_update_fields_on_open(doc)
        doc.save(str(out))
        _log(f"已保存 docx: {out}")
        # 在 save 之后调用 Word COM 离线刷新所有域（TOC / PAGE / REF），
        # 让用户拿到的 docx 自带真实目录与页码。无 Word 时回退到上面的开关模式。
        try:
            self._update_fields_via_word(out)
        except Exception as e:
            _log(f"COM 刷新域失败（非致命）: {e}", "WARNING")
        # 锁定封面页（在 COM 刷新域之后，避免 Word 重新保存时清除 permStart/permEnd）：
        # 重新加载磁盘 docx（COM 可能已经重新保存），写入 documentProtection + permStart/permEnd
        try:
            from docx import Document as _Document
            doc2 = _Document(str(out))
            # 封面页保护已取消（用户要求封面可编辑），不再调用 _lock_cover_page
            doc2.save(str(out))
            _log("封面页保护已取消（封面可编辑）")
        except Exception as e:
            _log(f"更新域流程异常（docx 已保存可用）: {e}", "WARNING")
        return out

    @staticmethod
    def _enable_update_fields_on_open(doc):
        """在 settings.xml 中写入 <w:updateFields w:val="true"/>，让 Word 打开
        文档时自动更新所有域（TOC 等动态字段）。

        关键：CT_Settings schema 对子元素顺序有严格要求，<w:updateFields> 必须位于
        <w:rsids> 之前、<w:compat> 之后（具体顺序参见 ECMA-376 第 17.15.1.75 节
        附近的 schema 序列）。直接 append 到末尾会被某些 Word 版本忽略。
        这里按正确顺序插入。

        Word 行为：检测到该开关后，打开文档时弹"此文档包含可能引用其他文件的
        字段。是否更新此文档中的字段？"提示，点"是"即更新目录。
        """
        settings = doc.settings.element
        # 避免重复添加
        existing = settings.find(qn("w:updateFields"))
        if existing is not None:
            existing.set(qn("w:val"), "true")
            return
        upd = OxmlElement("w:updateFields")
        upd.set(qn("w:val"), "true")
        # 按 CT_Settings 顺序，updateFields 应在 rsids 之前。找 rsids 锚点插入
        rsids = settings.find(qn("w:rsids"))
        if rsids is not None:
            rsids.addprevious(upd)
            return
        # 回退：找 compat 锚点之后插入
        compat = settings.find(qn("w:compat"))
        if compat is not None:
            compat.addnext(upd)
            return
        # 最终回退：append 到末尾
        settings.append(upd)

    def _lock_cover_page(self, doc):
        """锁定封面页（不让编辑）：通过 documentProtection + range permission 实现。

        实现策略：
        1. 在 settings.xml 写入 <w:documentProtection w:edit="readOnly"
           w:enforcement="1"/>，让整个文档默认只读
        2. 找到封面 section 的 section break（在封面最后一段的 pPr.sectPr 里），
           在它之后插入 <w:permStart w:id="1" w:edGrp="everyone"/>，
           让正文区域可编辑
        3. 在 body 末尾（最后一个 sectPr 之前）插入 <w:permEnd w:id="1"/>

        用户可在 Word"审阅→限制编辑→停止保护"解除（无需密码）。
        """
        try:
            # 1) settings.xml 加 documentProtection readOnly
            settings = doc.settings.element
            existing = settings.find(qn("w:documentProtection"))
            if existing is not None:
                settings.remove(existing)
            dp = OxmlElement("w:documentProtection")
            dp.set(qn("w:edit"), "readOnly")
            dp.set(qn("w:enforcement"), "1")
            upd = settings.find(qn("w:updateFields"))
            if upd is not None:
                upd.addprevious(dp)
            else:
                rsids = settings.find(qn("w:rsids"))
                if rsids is not None:
                    rsids.addprevious(dp)
                else:
                    settings.append(dp)

            # 2) 找封面 section break（中间 section 的 sectPr 在段落 pPr 里）
            body = doc.element.body
            # 列出所有 sectPr（包括 body 直接的 + 段落 pPr 里的）
            all_sect_prs = body.findall('.//' + qn('w:sectPr'))
            if len(all_sect_prs) >= 2:
                # 第一个 sectPr 是封面 section break（在封面最后一段 pPr 里）
                cover_sect_pr = all_sect_prs[0]
                # 在封面 sectPr 所在段落之后插入 permStart
                # 找到包含 cover_sect_pr 的段落
                parent_p = cover_sect_pr.getparent()  # pPr
                parent_p = parent_p.getparent() if parent_p is not None else None  # p
                if parent_p is not None:
                    perm_start = OxmlElement("w:permStart")
                    perm_start.set(qn("w:id"), "1")
                    perm_start.set(qn("w:edGrp"), "everyone")
                    parent_p.addnext(perm_start)

                    # 在 body 末尾（最后 sectPr 之前）插入 permEnd
                    last_sect_pr = all_sect_prs[-1]
                    perm_end = OxmlElement("w:permEnd")
                    perm_end.set(qn("w:id"), "1")
                    last_sect_pr.addprevious(perm_end)
            elif len(all_sect_prs) == 1:
                # 只有一个 section（无封面分节），跳过锁定
                _log("仅 1 个 section，跳过封面锁定", "WARNING")
        except Exception as e:
            _log(f"封面页锁定失败: {e}", "WARNING")

    def _update_fields_via_word(self, out_path):
        """在 doc.save() 之后调用 Word COM 离线刷新所有域（TOC / PAGE / REF 等）。

        读取 config["update_fields_after_save"]：
          - enabled: 是否启用
          - backend: win32com（MS Word） / wps（Kwps.Application） / none
          - timeout_sec: COM 调用整体超时
          - fail_silently: 失败时仅告警不抛

        关键点：
          1. COM 调用必须 try/except/finally，确保 word.Quit() 一定执行，
             否则 Word 进程会泄漏卡死后续自动化。
          2. Word 单实例 + 文件被占用时会冲突，调用前用 os.open 独占尝试检测；
             被占用则跳过 COM，回退到 settings.xml 开关模式。
          3. 更新顺序：doc.Fields.Update → 遍历 TablesOfContents 逐个 Update →
             doc.Repaginate → doc.Save，确保页码与目录条目都已写入磁盘。
        """
        cfg = self.config.get("update_fields_after_save") or {}
        if not cfg.get("enabled", False):
            _log("update_fields_after_save.enabled=false，跳过 COM 更新域")
            return False
        backend = (cfg.get("backend") or "none").lower()
        if backend == "none":
            return False
        fail_silently = bool(cfg.get("fail_silently", True))
        timeout_sec = int(cfg.get("timeout_sec", 60) or 60)

        # 文件占用检测：用 os.open 以共享写方式探测，被独占则跳过 COM。
        # 注意：不能用 O_EXCL（文件已存在时本就返回 EEXIST，逻辑反了）。
        # Windows 上 os.open 默认不共享，被 Word/WPS 持有时会抛 PermissionError。
        try:
            fd = os.open(str(out_path), os.O_RDWR)
            os.close(fd)
        except OSError:
            _log(f"docx 已被占用，跳过 COM 更新域: {out_path}", "WARNING")
            return False

        # 迟延导入，避免无 pywin32 环境直接 ImportError
        try:
            import win32com.client as _w32c  # noqa: F401
            import pythoncom  # noqa: F401
        except ImportError:
            msg = "pywin32 未安装，无法调用 Word COM 更新域；回退到 settings.xml 开关模式"
            if fail_silently:
                _log(msg, "WARNING")
                return False
            raise

        # 依序尝试候选 ProgID：win32com 后端先试 MS Word，失败再降到 WPS；
        # 只装 WPS 的机器没有 Word.Application，直接抛 CO_E_CLASSSTRING（无效类字符串）。
        if backend == "win32com":
            progid_candidates = ("Word.Application", "Kwps.Application")
        else:
            progid_candidates = ("Kwps.Application",)
        word = None
        doc = None
        last_err = None
        for progid in progid_candidates:
            try:
                pythoncom.CoInitialize()
                word = _w32c.DispatchEx(progid)
                try:
                    word.Visible = False
                except Exception:
                    pass
                try:
                    word.DisplayAlerts = 0  # wdAlertsNone
                except Exception:
                    pass
                break
            except Exception as e:
                last_err = e
                word = None
                try:
                    pythoncom.CoUninitialize()
                except Exception:
                    pass
        try:
            if word is None:
                raise last_err if last_err is not None else RuntimeError("无可用 COM 后端")

            # Word 2007 等老版本对相对路径 / 正斜杠的 Documents.Open 不稳定，
            # 会触发 Office 文件打开子组件调用系统注册的 Shell Hook（WPS 安装残留
            # 仍可能注册），导致异常 source 标为 'Kingsoft WPS'、错误码 3010
            # '文档打开失败'。强制转绝对路径 + 反斜杠避开该路径。
            abs_win_path = os.path.abspath(str(out_path)).replace("/", "\\")
            doc = word.Documents.Open(abs_win_path, ReadOnly=False)

            # 更新所有域
            try:
                doc.Fields.Update()
            except Exception as e:
                _log(f"Fields.Update 异常: {e}", "WARNING")

            # 优先更新 TOC（如果存在）
            toc_count = 0
            try:
                tocs = doc.TablesOfContents
                toc_count = tocs.Count
                for i in range(1, toc_count + 1):
                    try:
                        tocs(i).Update()
                    except Exception as e:
                        _log(f"TOC[{i}].Update 异常: {e}", "WARNING")
            except Exception as e:
                _log(f"TablesOfContents 访问异常: {e}", "WARNING")

            # 重新分页确保页码刷新
            try:
                doc.Repaginate()
            except Exception as e:
                _log(f"Repaginate 异常: {e}", "WARNING")

            # 保存（用原格式，避免触发 SaveAs 转换对话框）
            try:
                doc.Save()
            except Exception as e:
                _log(f"doc.Save 异常: {e}", "WARNING")

            _log(f"Word COM 已更新域（TOC: {toc_count}）: {out_path}")
            return True
        except Exception as e:
            msg = f"Word COM 更新域失败 [{backend}]: {e}"
            if fail_silently:
                _log(msg, "WARNING")
                return False
            raise
        finally:
            try:
                if doc is not None:
                    doc.Close(SaveChanges=0)
            except Exception:
                pass
            try:
                if word is not None:
                    word.Quit()
            except Exception:
                pass
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass


    def _insert_toc(self, doc, levels="1-3"):
        """在文档开头插入目录（TOC 域），标题"目录"为普通段落。

        生成的结构：
          段落1：标题"目录"（居中、加粗、Heading 1 样式以方便识别）
          段落2：TOC 域（fldChar begin → instrText 'TOC \\o "1-3" \\h \\z' →
                  fldChar separate → 占位文本 → fldChar end）
        Word 打开后会提示"更新域"，或用户右键→更新域，即可填充真实目录。
        """
        # 标题段
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run("目录")
        title_run.bold = True
        title_run.font.size = Pt(18)

        # TOC 域段
        toc_p = doc.add_paragraph()
        # 域开始：标记 w:dirty="true"，让 Word 打开时自动重算该域。
        run_begin = toc_p.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        fld_begin.set(qn("w:dirty"), "true")
        run_begin._element.append(fld_begin)

        # 域指令
        run_instr = toc_p.add_run()
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f' TOC \\o "{levels}" \\h \\z '
        run_instr._element.append(instr)

        # 域分隔
        run_sep = toc_p.add_run()
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        run_sep._element.append(fld_sep)

        # 占位文本（Word 打开后会被替换）
        run_placeholder = toc_p.add_run("（右键此处选择“更新域”以生成目录）")
        run_placeholder.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        # 域结束
        run_end = toc_p.add_run()
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run_end._element.append(fld_end)

        # 在目录后插入分页符，让正文从新页开始
        page_break_p = doc.add_paragraph()
        pb_run = page_break_p.add_run()
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        pb_run._element.append(br)

    def _insert_doc_info_page(self, doc):
        """文档信息页已合并到封面页（_append_doc_info_into_cover），此处保留为
        no-op 以避免在 TOC 前再生成一次重复内容。"""
        return

    def _append_doc_info_into_cover(self, doc, with_indent=False, top_padding_mm=0):
        """追加版权申明 + 文档信息。

        若 with_indent=True（封面 section 边距为 0 时的第 2 页场景），
        为所有段落加 10mm 左右缩进，模拟 A4 标准内边距视觉效果。
        若 top_padding_mm>0，给第一个段落（h2）额外加 space_before，
        让内容离页眉保留距离。
        """
        node = getattr(self, "_doc_info_node", None)
        if node is None:
            _log("未找到 .sr-doc-info-page 节点，跳过封面文档信息追加", "WARNING")
            return
        fonts_cfg = self.config.get("fonts", {}) or {}
        container = _Container(doc, fonts=fonts_cfg)

        def _apply_indent(p):
            if with_indent:
                pf = p.paragraph_format
                pf.left_indent = Mm(10)
                pf.right_indent = Mm(10)

        first_block = True
        for child in node.children:
            if child.name in ('h2', 'h3'):
                text = child.get_text(" ", strip=True)
                if not text:
                    continue
                is_h2 = child.name == 'h2'
                p = container.add_paragraph("")
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                pf = p.paragraph_format
                # 用户调整后的间距：h2 3.5/2.1mm, h3 2.8/1.4mm
                base_before = 3.5 if is_h2 else 2.8
                if first_block and top_padding_mm > 0:
                    # 第一个 h2 加额外 top_padding
                    base_before += top_padding_mm
                pf.space_before = Mm(base_before)
                pf.space_after = Mm(2.1 if is_h2 else 1.4)
                _apply_indent(p)
                # 标题前加无序项目符号 "▪ "（方形项目符号）
                run = p.add_run("▪  ")
                run.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                first_block = False
            elif child.name == 'p':
                text = child.get_text(" ", strip=True)
                if not text:
                    continue
                p = container.add_paragraph("")
                _apply_indent(p)
                # 首行缩进 7.4mm（2em 中文段首）
                p.paragraph_format.first_line_indent = Mm(7.4)
                # 加大行距：1.5 倍
                p.paragraph_format.line_spacing = 1.5
                run = p.add_run(text)
                run.font.name = '微软雅黑'
                run.font.size = Pt(12)
                rpr = run._element.get_or_add_rPr()
                rFonts = rpr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rpr.insert(0, rFonts)
                rFonts.set(qn('w:eastAsia'), '微软雅黑')
                first_block = False
            elif child.name == 'table':
                self._map_node(child, container)
                last_tbl = doc.tables[-1] if doc.tables else None
                if last_tbl is not None:
                    tbl_pr = last_tbl._element.find(qn('w:tblPr'))
                    if tbl_pr is not None:
                        existing_jc = tbl_pr.find(qn('w:jc'))
                        if existing_jc is not None:
                            tbl_pr.remove(existing_jc)
                        jc = OxmlElement('w:jc')
                        jc.set(qn('w:val'), 'center')
                        tbl_pr.append(jc)
                        if with_indent:
                            tbl_ind = tbl_pr.find(qn('w:tblInd'))
                            if tbl_ind is None:
                                tbl_ind = OxmlElement('w:tblInd')
                                tbl_pr.append(tbl_ind)
                            tbl_ind.set(qn('w:w'), str(int(Mm(10).twips)))
                            tbl_ind.set(qn('w:type'), 'dxa')
                    # 增大单元格行距：给所有单元格内的段落加 1.5 倍行距
                    # 文档信息表（sr-copyright-meta）除外：1.5 倍行距会把行高撑到
                    # 24pt 而文本只有 13.8pt，Word 的 vAlign=center 只能居中"段落框"，
                    # 行距的 leading 又压在字形上方，导致内容视觉偏上。此处改为
                    # 单倍行距 + 最小段距，让行贴近文本后由 vAlign=center 真正居中。
                    is_doc_info = "sr-copyright-meta" in (child.get("class") or [])
                    for row in last_tbl.rows:
                        for cell in row.cells:
                            for cell_p in cell.paragraphs:
                                if is_doc_info:
                                    cell_p.paragraph_format.line_spacing = 1.0
                                    cell_p.paragraph_format.space_before = Pt(0)
                                    cell_p.paragraph_format.space_after = Pt(0)
                                else:
                                    cell_p.paragraph_format.line_spacing = 1.5
                                    cell_p.paragraph_format.space_before = Pt(2)
                                    cell_p.paragraph_format.space_after = Pt(2)
                first_block = False

    # ── 首页（封面）──────────────────────────────────

    def _extract_cover_info(self, root):
        """从 DOM root 中提取封面信息（标题、客户名、报告时间）。

        HTML 结构：
          <header class="sr-cover" id="report-top">
            ...
            <h1 class="report-cover-title-main" data-field="cover.title">...</h1>
            ...
            <span class="sr-cover-meta-label">客户名称</span>
            <span class="report-cover-title-sub" data-field="cover.clientName">...</span>
            ...
            <span class="sr-cover-meta-label">报告时间</span>
            <span class="report-cover-desc" data-field="cover.period">...</span>
          </header>
        返回 dict: {title, client_name, period}；找不到的项为空字符串。
        封面节点 .sr-cover 在 dom_root 之外，extract_dom 时已缓存到 self._cover_node。
        """
        info = {"title": "", "client_name": "", "period": ""}
        cover = getattr(self, "_cover_node", None) or (root.select_one(".sr-cover") if root else None)
        if cover is None:
            return info

        def _text(sel):
            el = cover.select_one(sel)
            return el.get_text(strip=True) if el else ""

        info["title"] = _text(".report-cover-title-main")
        info["client_name"] = _text(".report-cover-title-sub")
        info["period"] = _text(".report-cover-desc")
        return info

    def _cover_bg_path(self):
        """封面背景图固定路径：html_to_word/assets/a4-portrait-bg.png"""
        return Path(__file__).resolve().parent / "assets" / "a4-portrait-bg.png"

    def _title_page_bg_path(self):
        """标题页背景图（含背景+logo+客户名/日期渐变色块+底部曲线）：
        html_to_word/title-page-bg.png，1200x544 px 横向图。"""
        return Path(__file__).resolve().parent / "title-page-bg.png"

    def _insert_cover_page(self, doc, cover_info):
        """封面页：按 HTML A4 竖版封面布局。

        布局（A4 竖版 210×297mm，对应 Figma 1200×1700）：
          第 1 页（封面）：
            0~297mm    : title-page-bg.png 整页背景（1200×1700 RGB）+ logo（上方居中）
            ~239mm     : 报告标题（bottom: 330/1700 ≈ 57.5mm 处，font-size 11.18mm）
            ~261mm     : 客户名称（bottom: 206/1700 ≈ 36mm 处，font-size 3.85mm）
            ~265mm     : 报告日期（紧接客户名，font-size 3.5mm）
          第 2 页（版权/文档信息）：
            版权申明 h2 + 段落 + 文档信息 h3 + 表格

        实现：title-page-bg.png 作为衬于文字下方的浮动图片，覆盖整页（210×297mm）；
        标题/客户名/日期段落通过精确 space_before 推到 Figma 对应位置；
        末尾插入分页符，让版权/文档信息落在新页。
        """
        title_bg_path = self._title_page_bg_path()
        if not title_bg_path.exists():
            _log(f"封面 title-page-bg.png 不存在: {title_bg_path}", "WARNING")
        section = doc.sections[0]
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        # 上下左右 0 边距，让背景图覆盖整页；后续段落通过 space_before 精确定位
        section.left_margin = Mm(0)
        section.right_margin = Mm(0)
        section.top_margin = Mm(0)
        section.bottom_margin = Mm(0)

        # 1) 插入 title-page-bg.png 浮动图片（衬于文字下方，定位到页面 (0,0)，覆盖整页 210×297mm）
        if title_bg_path.exists():
            try:
                self._add_floating_title_bg(doc, str(title_bg_path),
                                            page_w_mm=210, page_h_mm=297)
            except Exception as e:
                _log(f"title-page-bg.png 插入失败: {e}", "WARNING")
            # 给图片占位段加 line=48pt 行距，撑高该段约 17mm，让后续段落往下推
            last_p = doc.paragraphs[-1]
            last_p.paragraph_format.line_spacing = Pt(48)
            pPr = last_p._element.get_or_add_pPr()
            spacing = pPr.find(qn('w:spacing'))
            if spacing is not None:
                spacing.set(qn('w:lineRule'), 'auto')

        # 2) 报告标题：26pt 加粗 黑色 居中
        #    HTML A4 竖版 Figma 中 bottom:330/1700 ≈ 57.5mm（即距顶 239.5mm）；
        #    但 title-page-bg.png 在 y=1280 (~223mm) 以下为纯白区域，
        #    用户反馈文字离页脚太近，调整到非白区域中下部：~180mm 处。
        title = cover_info.get("title", "").strip()
        if title:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Mm(155)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(title)
            run.bold = True
            run.font.size = Pt(26)
            run.font.color.rgb = RGBColor(0x1E, 0x23, 0x2B)

        # 3) 客户名称：14pt 加粗 #2F3540 居中
        #    跟随标题，目标位置 ~198mm
        client = cover_info.get("client_name", "").strip()
        if client:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Mm(7)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(client)
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x2F, 0x35, 0x40)

        # 4) 报告日期：13pt #2F3540 居中（紧接客户名）
        period = cover_info.get("period", "").strip()
        if period:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Mm(3)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(period)
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x2F, 0x35, 0x40)

        # 5) 插入分页符：让版权申明 + 文档信息落到第 2 页
        page_break_p = doc.add_paragraph()
        page_break_p.add_run().add_break(WD_BREAK.PAGE)

        # 6) 第 2 页顶部加 25mm 空 spacer 段，把版权申明从页眉处推开
        #    （直接用 space_before 容易被分页符后的段间距折叠，空段更可靠）
        top_spacer = doc.add_paragraph()
        top_spacer.paragraph_format.space_before = Pt(0)
        top_spacer.paragraph_format.space_after = Pt(0)
        top_spacer.paragraph_format.line_spacing = Pt(50)  # 50pt ≈ 17.6mm
        pPr_sp = top_spacer._element.get_or_add_pPr()
        spacing_sp = pPr_sp.find(qn('w:spacing'))
        if spacing_sp is not None:
            spacing_sp.set(qn('w:lineRule'), 'exact')

        # 7) 版权申明 + 文档信息（在第 2 页，距离顶部约 18mm）
        #    第 2 页仍在封面 section（0 边距），左右通过 with_indent 加 10mm 缩进
        self._append_doc_info_into_cover(doc, with_indent=True)

    def _add_floating_title_bg(self, doc, image_path, page_w_mm, page_h_mm):
        """插入 title-page-bg.png 作为浮动图片：behindDoc=1（衬于文字下方），
        绝对定位到页面 (0,0)，宽度撑满 page_w_mm（210mm），高度按图片自身
        像素比例缩放（1200x1700 → 210mm x 297mm，正好覆盖整页）。

        关键：直接强制 width=210mm，按像素比例计算 height，**不**用 96 DPI
        换算（避免图片实际无 DPI 信息时被算成 317mm 然后错误缩放成 21mm）。
        """
        from PIL import Image as _PILImage
        with _PILImage.open(image_path) as im:
            px_w, px_h = im.size
        # 按像素比例计算高度
        final_w_mm = page_w_mm
        final_h_mm = page_w_mm * px_h / px_w  # 210 * 1700/1200 ≈ 297.5mm（整页）
        # 如果计算出的高度超过 page_h_mm，截到 page_h_mm
        if final_h_mm > page_h_mm:
            final_h_mm = page_h_mm
        final_w_emu = int(Mm(final_w_mm).emu)
        final_h_emu = int(Mm(final_h_mm).emu)

        # 1) 用 add_picture 插入内嵌图片，再转为 anchor
        pic_para = doc.add_paragraph()
        pf = pic_para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        run = pic_para.add_run()
        run.add_picture(image_path, width=Mm(final_w_mm), height=Mm(final_h_mm))

        # 2) 找 <wp:inline>，转 <wp:anchor>
        r = run._element
        drawing = r.find(qn("w:drawing"))
        if drawing is None:
            _log("未找到 w:drawing，title-page-bg 插入回退", "WARNING")
            return
        inline = drawing.find(qn("wp:inline"))
        if inline is None:
            _log("未找到 wp:inline，title-page-bg 插入回退", "WARNING")
            return

        anchor = OxmlElement("wp:anchor")
        anchor.set("distT", "0")
        anchor.set("distB", "0")
        anchor.set("distL", "0")
        anchor.set("distR", "0")
        anchor.set("simplePos", "0")
        anchor.set("relativeHeight", "0")
        anchor.set("behindDoc", "1")
        anchor.set("locked", "0")
        anchor.set("layoutInCell", "1")
        anchor.set("allowOverlap", "1")

        simple_pos = OxmlElement("wp:simplePos")
        simple_pos.set("x", "0")
        simple_pos.set("y", "0")
        anchor.append(simple_pos)

        pos_h = OxmlElement("wp:positionH")
        pos_h.set("relativeFrom", "page")
        pos_offset = OxmlElement("wp:posOffset")
        pos_offset.text = "0"
        pos_h.append(pos_offset)
        anchor.append(pos_h)

        pos_v = OxmlElement("wp:positionV")
        pos_v.set("relativeFrom", "page")
        pos_offset_v = OxmlElement("wp:posOffset")
        pos_offset_v.text = "0"
        pos_v.append(pos_offset_v)
        anchor.append(pos_v)

        extent = OxmlElement("wp:extent")
        extent.set("cx", str(final_w_emu))
        extent.set("cy", str(final_h_emu))
        anchor.append(extent)

        effect_extent = OxmlElement("wp:effectExtent")
        for side in ('l', 't', 'r', 'b'):
            effect_extent.set(side, "0")
        anchor.append(effect_extent)

        wrap_none = OxmlElement("wp:wrapNone")
        anchor.append(wrap_none)

        doc_pr = OxmlElement("wp:docPr")
        doc_pr.set("id", "0")
        doc_pr.set("name", "封面标题图")
        anchor.append(doc_pr)

        cnv = OxmlElement("wp:cNvGraphicFramePr")
        anchor.append(cnv)

        graphic = inline.find(qn("a:graphic"))
        if graphic is not None:
            anchor.append(graphic)

        drawing.remove(inline)
        drawing.append(anchor)

    def _add_floating_background_picture(self, doc, image_path, page_w_mm, page_h_mm):
        """插入一张浮动图片：behindDoc=1（衬于文字下方），绝对定位到页面
        (0,0)，extent=整页大小。后续段落流式排列但视觉上叠加在图片之上。

        python-docx 的 add_picture 默认生成 <wp:inline>，需要后处理为
        <wp:anchor>。做法：先用 add_picture 创建 inline 图片段落，再读取
        该段落的 <wp:inline>，转换为 <wp:anchor> 并补充 behindDoc/positionH/
        positionV/extent/simplePos 子元素。
        """
        from PIL import Image as _PILImage
        # 计算图片实际比例，按"完整覆盖页面"等比缩放
        with _PILImage.open(image_path) as im:
            px_w, px_h = im.size
        img_w_mm = px_w * 25.4 / 96.0
        img_h_mm = px_h * 25.4 / 96.0
        scale = max(page_w_mm / img_w_mm, page_h_mm / img_h_mm)
        final_w_emu = int(Mm(img_w_mm * scale).emu)
        final_h_emu = int(Mm(img_h_mm * scale).emu)

        # 1) 先用 add_picture 插入内嵌图片，获得包含 <wp:inline> 的段落
        pic_para = doc.add_paragraph()
        run = pic_para.add_run()
        run.add_picture(image_path, width=Mm(img_w_mm * scale), height=Mm(img_h_mm * scale))

        # 2) 找到 <wp:inline> 元素，转换为 <wp:anchor>
        r = run._element
        drawing = r.find(qn("w:drawing"))
        if drawing is None:
            _log("未找到 w:drawing，背景图插入回退", "WARNING")
            return
        inline = drawing.find(qn("wp:inline"))
        if inline is None:
            _log("未找到 wp:inline，背景图插入回退", "WARNING")
            return

        # 构建 <wp:anchor> 元素，设置 behindDoc=1，把图片衬于文字下方
        nsmap_qn = qn
        anchor = OxmlElement("wp:anchor")
        anchor.set("distT", "0")
        anchor.set("distB", "0")
        anchor.set("distL", "0")
        anchor.set("distR", "0")
        anchor.set("simplePos", "0")
        anchor.set("relativeHeight", "0")
        anchor.set("behindDoc", "1")
        anchor.set("locked", "0")
        anchor.set("layoutInCell", "1")
        anchor.set("allowOverlap", "1")

        # simplePos 元素（必须存在但 simplePos=0 时被忽略）
        simple_pos = OxmlElement("wp:simplePos")
        simple_pos.set("x", "0")
        simple_pos.set("y", "0")
        anchor.append(simple_pos)

        # positionH：相对 page，偏移 0
        pos_h = OxmlElement("wp:positionH")
        pos_h.set("relativeFrom", "page")
        pos_offset = OxmlElement("wp:posOffset")
        pos_offset.text = "0"
        pos_h.append(pos_offset)
        anchor.append(pos_h)

        # positionV：相对 page，偏移 0
        pos_v = OxmlElement("wp:positionV")
        pos_v.set("relativeFrom", "page")
        pos_offset_v = OxmlElement("wp:posOffset")
        pos_offset_v.text = "0"
        pos_v.append(pos_offset_v)
        anchor.append(pos_v)

        # extent：图片最终大小
        extent = OxmlElement("wp:extent")
        extent.set("cx", str(final_w_emu))
        extent.set("cy", str(final_h_emu))
        anchor.append(extent)

        # effectExtent（占位）
        effect_extent = OxmlElement("wp:effectExtent")
        effect_extent.set("l", "0")
        effect_extent.set("t", "0")
        effect_extent.set("r", "0")
        effect_extent.set("b", "0")
        anchor.append(effect_extent)

        # wrapNone：让文字穿过图片（图片不占用段落空间）
        wrap_none = OxmlElement("wp:wrapNone")
        anchor.append(wrap_none)

        # docPr
        doc_pr = OxmlElement("wp:docPr")
        doc_pr.set("id", "0")
        doc_pr.set("name", "背景图")
        anchor.append(doc_pr)

        # cNvGraphicFramePr（占位）
        cnv = OxmlElement("wp:cNvGraphicFramePr")
        anchor.append(cnv)

        # 把原 <wp:inline> 的 <a:graphic> 子元素搬到 <wp:anchor> 下
        graphic = inline.find(qn("a:graphic"))
        if graphic is not None:
            anchor.append(graphic)

        # 替换 inline 为 anchor
        drawing.remove(inline)
        drawing.append(anchor)

        # 让该段落自身不占用额外空间（行高最小）
        pf = pic_para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0

    @staticmethod
    def _add_cover_spacer(doc, mm_before):
        """插入空白段，用前置间距把后续内容往下推 mm_before 毫米。"""
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        # 用前置间距实现下移
        from docx.shared import Mm as _Mm
        pf.space_before = _Mm(mm_before)
        return p

    def _add_white_card_panel(self, doc, width_mm=190, padding_mm=3,
                              border_hex="D3D7DE", indent_left_mm=0,
                              fill_hex="FFFFFF", align="center"):
        """创建一个 1×1 卡片容器表格，用于封面里的"卡片式"信息块。

        - fill_hex 底色（默认白 FFFFFF；可用 F5F7FA 等浅色模拟半透明叠加效果）
        - 浅灰边框（border_hex，1pt）+ 固定宽度
        - align: 表格水平对齐方式（center / left / right），默认居中
          居中时 tblInd=0 + tblPrJc=center，避免硬左缩导致整体偏右
        - 上下左右内边距 padding_mm
        - fixed 布局，nil 默认样式

        返回单元格对象（_Cell），调用方可往里 add_paragraph / add_table。
        Word 表格本身不支持圆角；这里用浅灰边框 + 白底 + 内边距模拟卡片视觉。
        """
        tbl = doc.add_table(rows=1, cols=1)
        tbl_pr = tbl._element.find(qn('w:tblPr'))
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            tbl._element.insert(0, tbl_pr)
        # 清掉 python-docx 默认 tblStyle / 边框
        for tag in ('w:tblStyle', 'w:tblBorders', 'w:tblW', 'w:tblInd',
                    'w:tblCellMar', 'w:tblLayout', 'w:tblPrJc'):
            el = tbl_pr.find(qn(tag))
            if el is not None:
                tbl_pr.remove(el)
        # tblW 固定宽度
        tbl_w = OxmlElement('w:tblW')
        tbl_w.set(qn('w:type'), 'dxa')
        tbl_w.set(qn('w:w'), str(int(Mm(width_mm).twips)))
        tbl_pr.append(tbl_w)
        # 对齐方式：center 用 tblPrJc；left 用 tblInd 左缩进
        if align == "center":
            tbl_jc = OxmlElement('w:jc')
            tbl_jc.set(qn('w:val'), 'center')
            tbl_pr.append(tbl_jc)
            # 居中时 tblInd=0
            tbl_ind = OxmlElement('w:tblInd')
            tbl_ind.set(qn('w:type'), 'dxa')
            tbl_ind.set(qn('w:w'), '0')
            tbl_pr.append(tbl_ind)
        else:
            tbl_ind = OxmlElement('w:tblInd')
            tbl_ind.set(qn('w:type'), 'dxa')
            tbl_ind.set(qn('w:w'), str(int(Mm(indent_left_mm).twips)))
            tbl_pr.append(tbl_ind)
        # tblLayout fixed
        tbl_layout = OxmlElement('w:tblLayout')
        tbl_layout.set(qn('w:type'), 'fixed')
        tbl_pr.append(tbl_layout)
        # 边框：单线，1pt = 8 twips, 颜色 border_hex
        tbl_borders = OxmlElement('w:tblBorders')
        for side in ('top', 'left', 'bottom', 'right'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '8')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), border_hex)
            tbl_borders.append(b)
        for side in ('insideH', 'insideV'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'nil')
            tbl_borders.append(b)
        tbl_pr.append(tbl_borders)
        # 内边距
        pad_twips = int(Mm(padding_mm).twips)
        tbl_cell_mar = OxmlElement('w:tblCellMar')
        for side in ('top', 'left', 'bottom', 'right'):
            m = OxmlElement(f'w:{side}')
            m.set(qn('w:w'), str(pad_twips))
            m.set(qn('w:type'), 'dxa')
            tbl_cell_mar.append(m)
        tbl_pr.append(tbl_cell_mar)
        # 底色
        cell = tbl.rows[0].cells[0]
        self._set_cell_shading(cell, fill_hex)
        # 单元格宽度 = width - 2 * padding
        cell.width = Mm(max(width_mm - 2 * padding_mm, 10))
        return cell

    def _switch_to_body_section(self, doc):
        """在封面后新增 next-page section，作为正文 section（A4 竖版，与封面统一）。

        python-docx: doc.add_section(WD_SECTION.NEW_PAGE) 会插入分节符并新增 section。
        之后把新 section 的页面尺寸/边距重新设为正文（A4 竖版），并清掉新 section
        继承的页眉页脚链接以避免影响正文。
        """
        from docx.enum.section import WD_SECTION
        new_section = doc.add_section(WD_SECTION.NEW_PAGE)
        cfg = self.config.get("docx", {})
        w = Mm(cfg.get("page_width_mm", 210))
        h = Mm(cfg.get("page_height_mm", 297))
        m = Mm(cfg.get("margin_mm", 20))
        new_section.orientation = WD_ORIENT.PORTRAIT
        new_section.page_width = w
        new_section.page_height = h
        new_section.left_margin = m
        new_section.right_margin = m
        new_section.top_margin = m
        new_section.bottom_margin = m
        # 给正文 section 加居中页码页脚（封面 section 不链接，保持无页码）
        self._add_page_number_footer(new_section)
        # 给正文 section 加页眉图片（封面 section 不链接）
        self._add_header_image(new_section)
        return new_section

    def _add_header_image(self, section):
        """给指定 section 的 header 加页眉图片（PNG）。

        配置项 header_image.path 为空时跳过。
        align: left / center / right，默认 left。
        """
        cfg = self.config.get("header_image") or {}
        path = (cfg.get("path") or "").strip()
        if not path:
            return
        from pathlib import Path as _Path
        img_path = _Path(path)
        if not img_path.is_absolute():
            img_path = _Path(__file__).resolve().parent / img_path
        if not img_path.exists():
            _log(f"页眉图片文件不存在，跳过: {path}", "WARNING")
            return
        height_mm = float(cfg.get("height_mm") or 10)
        align_str = str(cfg.get("align") or "left").lower()
        align_map = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                     "center": WD_ALIGN_PARAGRAPH.CENTER,
                     "right": WD_ALIGN_PARAGRAPH.RIGHT}
        align = align_map.get(align_str, WD_ALIGN_PARAGRAPH.LEFT)

        header = section.header
        header.is_linked_to_previous = False
        section.header_distance = Mm(float(cfg.get("header_distance_mm") or 12.7))
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.clear()
        p.alignment = align
        run = p.add_run()
        run.add_picture(str(img_path), height=Mm(height_mm))

    def _add_page_number_footer(self, section):
        """给指定 section 的 footer 加居中页码（PAGE / NUMPAGES 域），
        格式：当前页 / 总页数。封面 section 不被调用。

        实现：section.footer.is_linked_to_previous = False 断开继承，
        在 footer 段落里插入 PAGE 域 + " / " + NUMPAGES 域，居中对齐。
        同时设置 pgNumType start=1，让正文 section 从 1 开始计数（封面不计入正文页码）。
        """
        try:
            section.footer.is_linked_to_previous = False
            footer = section.footer
            # 清空已有段落
            for p in list(footer.paragraphs):
                p._element.getparent().remove(p._element)
            p = footer.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            def _add_field(run_p, instr_text, placeholder="1", size_pt=10.5):
                """在 run_p 末尾插入一个 Word 域（fldChar/instrText/separate/占位/end）。"""
                run_begin = run_p.add_run()
                fld_begin = OxmlElement("w:fldChar")
                fld_begin.set(qn("w:fldCharType"), "begin")
                run_begin._element.append(fld_begin)

                run_instr = run_p.add_run()
                instr = OxmlElement("w:instrText")
                instr.set(qn("xml:space"), "preserve")
                instr.text = f" {instr_text} \\* MERGEFORMAT "
                run_instr._element.append(instr)

                run_sep = run_p.add_run()
                fld_sep = OxmlElement("w:fldChar")
                fld_sep.set(qn("w:fldCharType"), "separate")
                run_sep._element.append(fld_sep)

                run_placeholder = run_p.add_run(placeholder)
                run_placeholder.font.size = Pt(size_pt)

                run_end = run_p.add_run()
                fld_end = OxmlElement("w:fldChar")
                fld_end.set(qn("w:fldCharType"), "end")
                run_end._element.append(fld_end)

            # 插入 PAGE 域
            _add_field(p, "PAGE", "1")
            # 分隔符 " / "
            sep_run = p.add_run(" / ")
            sep_run.font.size = Pt(10.5)
            # 插入 SECTIONPAGES 域（只算当前 section 的总页数 = 正文 section 页数，
            # 不含封面/版权页）
            _add_field(p, "SECTIONPAGES", "1")

            # 让正文 section 页码从 1 开始（封面 section 不带页码，从节起始重算）
            sect_pr = section._sectPr
            existing_pgnum = sect_pr.find(qn("w:pgNumType"))
            if existing_pgnum is not None:
                sect_pr.remove(existing_pgnum)
            pg_num = OxmlElement("w:pgNumType")
            pg_num.set(qn("w:start"), "1")
            sect_pr.append(pg_num)
        except Exception as e:
            _log(f"添加页码页脚失败: {e}", "WARNING")

    def _configure_page(self, doc):
        """设置 A4 竖版页面与边距。"""
        cfg = self.config.get("docx", {})
        w = Mm(cfg.get("page_width_mm", 210))
        h = Mm(cfg.get("page_height_mm", 297))
        m = Mm(cfg.get("margin_mm", 20))
        section = doc.sections[0]
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = w
        section.page_height = h
        section.left_margin = m
        section.right_margin = m
        section.top_margin = m
        section.bottom_margin = m
        # 正文默认字体（应用 fonts.normal 配置）
        fonts_cfg = self.config.get("fonts", {}) or {}
        self._apply_font_style(doc.styles["Normal"], fonts_cfg.get("normal", {}))
        # 关闭 Heading 1~6 的自动编号：HTML 中已带手写序号（"1." "3.1" "4.1.1"），
        # Word 内置 Heading 样式默认链接到多级列表 numId，会叠加出双序号。
        # 这里直接从样式定义里删除 numPr，并在每个 add_heading 调用后再做一次兜底
        # 清理（见 _strip_paragraph_num_pr）。
        # 同时应用 fonts.heading_1~6 配置（字体/字号/颜色/加粗/间距）。
        for lvl in range(1, 7):
            try:
                hs = doc.styles[f"Heading {lvl}"]
                self._strip_style_num_pr(hs)
                self._apply_font_style(hs, fonts_cfg.get(f"heading_{lvl}", {}))
            except KeyError:
                continue
        # Normal 也要兜底清掉 numPr，避免某些模板默认带 List Number 段落样式
        self._strip_style_num_pr(doc.styles["Normal"])

    @staticmethod
    def _apply_font_style(style, cfg):
        """把 fonts.* 配置段应用到 docx 样式对象。

        支持字段：
          - font:        拉丁字体名（style.font.name）
          - east_asia:   东亚字体名（w:eastAsia）
          - size_pt:     字号（磅）
          - bold:        是否加粗
          - italic:      是否斜体
          - color_hex:   字体颜色（RRGGBB，不带 #）
          - space_before_pt / space_after_pt: 段前/段后间距（磅）
          - line_spacing_pt: 行距（磅）
        缺省字段不覆盖，保留 docx 模板默认。
        """
        if not cfg:
            return
        font = style.font
        if "font" in cfg and cfg["font"]:
            font.name = cfg["font"]
        if "size_pt" in cfg and cfg["size_pt"] is not None:
            font.size = Pt(cfg["size_pt"])
        if "bold" in cfg and cfg["bold"] is not None:
            font.bold = cfg["bold"]
        if "italic" in cfg and cfg["italic"] is not None:
            font.italic = cfg["italic"]
        if "color_hex" in cfg and cfg["color_hex"]:
            try:
                font.color.rgb = RGBColor.from_string(cfg["color_hex"])
            except (ValueError, AttributeError):
                _log(f"无效的 color_hex: {cfg['color_hex']}", "WARNING")
        # 东亚字体：写入 rFonts/@w:eastAsia
        if "east_asia" in cfg and cfg["east_asia"]:
            rPr = style.element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            # 清除主题字体引用，避免 Word 渲染中文时 fallback 到主题字体
            # （python-docx 默认模板 majorFont/Jpan=MS Gothic、Hans=宋体，
            #  仅追加显式 w:eastAsia 不删 asciiTheme/eastAsiaTheme 等
            #  主题引用属性时，部分字符仍会被 Word 按 Unicode script 归到
            #  主题字体）。同时设置 ascii/hAnsi，让英文字符也走显式字体名。
            for theme_attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
                if rFonts.get(qn(f"w:{theme_attr}")) is not None:
                    del rFonts.attrib[qn(f"w:{theme_attr}")]
            rFonts.set(qn("w:eastAsia"), cfg["east_asia"])
            # 若配置同时提供拉丁字体名，则一并写入 ascii/hAnsi，避免英文 fallback
            if "font" in cfg and cfg["font"]:
                rFonts.set(qn("w:ascii"), cfg["font"])
                rFonts.set(qn("w:hAnsi"), cfg["font"])
        # 段落间距
        pf = style.paragraph_format
        if "space_before_pt" in cfg and cfg["space_before_pt"] is not None:
            pf.space_before = Pt(cfg["space_before_pt"])
        if "space_after_pt" in cfg and cfg["space_after_pt"] is not None:
            pf.space_after = Pt(cfg["space_after_pt"])
        if "line_spacing_pt" in cfg and cfg["line_spacing_pt"] is not None:
            pf.line_spacing = Pt(cfg["line_spacing_pt"])

    @staticmethod
    def _strip_style_num_pr(style):
        """从样式定义中移除 w:numPr，避免该样式下的段落继承自动编号。"""
        pPr = style.element.get_or_add_pPr()
        if pPr is None:
            return
        for tag in ("w:numPr",):
            el = pPr.find(qn(tag))
            if el is not None:
                pPr.remove(el)

    @staticmethod
    def _strip_paragraph_num_pr(paragraph):
        """清除单个段落上的自动编号（numPr），保留文本中已手写的序号。"""
        pPr = paragraph._p.get_or_add_pPr()
        for el in pPr.findall(qn("w:numPr")):
            pPr.remove(el)

    # ──────────────────────────────────────────────
    # DOM → docx 元素映射
    # ──────────────────────────────────────────────

    def _map_children(self, parent_tag, container, depth=0):
        """深度优先遍历 parent_tag 的直接子节点，映射到 docx。"""
        for child in parent_tag.children:
            self._map_node(child, container, depth)

    def _map_node(self, node, container, depth=0):
        """映射单个 DOM 节点到 docx。container 是 _Container。"""
        if isinstance(node, NavigableString):
            text = str(node).strip()
            if text:
                container.add_paragraph(text)
            return
        if not isinstance(node, Tag):
            return
        if getattr(node, "name", None) is None:
            return
        # 封面 .sr-cover 已在 _insert_cover_page 中作为独立 section 渲染，
        # 此处跳过避免正文流里重复出现。
        if node.name == "header" and "sr-cover" in (node.get("class") or []):
            return
        # 文档信息页 .sr-doc-info-page 已在 _insert_doc_info_page 中单独渲染，
        # 此处跳过避免正文流里重复出现。
        if node.name == "section" and "sr-doc-info-page" in (node.get("class") or []):
            return
        # 检查 class 匹配的 style_map 项（优先于标签）
        handler = self._lookup_style_handler(node)
        if handler:
            handler(node, container)
            return
        name = node.name
        if name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self._map_heading(node, container, level=int(name[1]))
        elif name == "p":
            self._map_paragraph(node, container)
        elif name == "table":
            self._map_table(node, container)
        elif name == "ul":
            self._map_list(node, container, ordered=False)
        elif name == "ol":
            self._map_list(node, container, ordered=True)
        elif name == "img":
            self._map_image(node, container)
        elif name in ("div", "section", "header", "article", "aside", "main", "footer", "figure", "figcaption"):
            # 容器节点：递归处理子节点
            self._map_children(node, container, depth + 1)
        elif name in ("span", "strong", "em", "b", "i", "a", "small", "sub", "sup", "label"):
            text = node.get_text(" ", strip=True)
            if text:
                p = container.add_paragraph()
                run = p.add_run(text)
                if name in ("strong", "b"):
                    run.bold = True
                elif name in ("em", "i"):
                    run.italic = True
        elif name in ("br", "hr"):
            container.add_paragraph("")
        else:
            text = node.get_text(" ", strip=True)
            if text:
                container.add_paragraph(text)

    def _lookup_style_handler(self, node):
        """根据 style_map 查找自定义 handler。命中类名优先。"""
        sm = self.config.get("style_map", {})
        classes = node.get("class") or []
        for cls in classes:
            key = f".{cls}"
            if key in sm:
                handler_name = sm[key]
                handler = getattr(self, f"_handle_{handler_name}", None)
                if handler:
                    return handler
        if node.name in sm:
            handler_name = sm[node.name]
            return getattr(self, f"_handle_{handler_name}", None)
        return None

    # ── 具体映射函数 ─────────────────────────────────

    def _handle_heading_1(self, node, container): self._map_heading(node, container, level=1)
    def _handle_heading_2(self, node, container): self._map_heading(node, container, level=2)
    def _handle_heading_3(self, node, container): self._map_heading(node, container, level=3)
    def _handle_heading_4(self, node, container): self._map_heading(node, container, level=4)
    def _handle_heading_5(self, node, container): self._map_heading(node, container, level=5)
    def _handle_heading_6(self, node, container): self._map_heading(node, container, level=6)
    def _handle_paragraph(self, node, container): self._map_paragraph(node, container)
    def _handle_table(self, node, container): self._map_table(node, container)
    def _handle_bullet_list(self, node, container): self._map_list(node, container, ordered=False)
    def _handle_number_list(self, node, container): self._map_list(node, container, ordered=True)
    def _handle_image(self, node, container): self._map_image(node, container)

    def _handle_risk_card_table(self, node, container):
        """风险卡：用单行单列表格 + 内部段落近似还原。

        用户 mark 调整：
          - head 前不要额外空段（reuse_first_para）
          - head 标题字号 > 11pt（13pt）
          - body 里 p 正文首行缩进 2em（≈ 7.4mm）
          - 步骤标题"1 网络防护"整体加缩进 5mm
          - 步骤下子列表整体加缩进 10mm
        """
        card_table = container.add_table(rows=1, cols=1)
        cell = card_table.cell(0, 0)
        cell_container = _Container(container.doc, cell=cell, fonts=container.fonts)
        # 标记当前在风险卡内（让 _map_paragraph / _map_risk_step 加缩进）
        self._in_risk_card = True
        # 用 cell 默认的第一个段落作为 head 起点，避免出现额外空段
        first_para_used = False
        for child in node.children:
            if isinstance(child, Tag):
                child_classes = child.get("class") or []
                is_head = (
                    child.name in ("h6", "h5", "h4")
                    or "sr-risk-card__head" in child_classes
                )
                if is_head:
                    # head：用 _handle_risk_card_head 处理，避免额外空段
                    self._handle_risk_card_head(child, cell_container,
                                                reuse_first_para=not first_para_used)
                    first_para_used = True
                else:
                    self._map_node(child, cell_container)
        # 离开风险卡
        self._in_risk_card = False
        self._set_table_borders(card_table)

    def _handle_kpi_paragraph(self, node, container):
        """KPI 卡：用加粗大字段落。"""
        text = node.get_text(" ", strip=True)
        if not text:
            return
        p = container.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)

    def _handle_tag_critical(self, node, container): self._map_tag(node, container, "tag_critical")
    def _handle_tag_high(self, node, container): self._map_tag(node, container, "tag_high")
    def _handle_tag_medium(self, node, container): self._map_tag(node, container, "tag_medium")
    def _handle_tag_medium_low(self, node, container): self._map_tag(node, container, "tag_medium_low")
    def _handle_tag_low(self, node, container): self._map_tag(node, container, "tag_low")
    def _handle_tag_info(self, node, container): self._map_tag(node, container, "tag_info")
    def _handle_tag_success(self, node, container): self._map_tag(node, container, "tag_success")
    def _handle_tag_blue(self, node, container): self._map_tag(node, container, "tag_blue")
    def _handle_tag_light(self, node, container): self._map_tag_light(node, container)

    def _map_tag_light(self, node, container):
        """浅底深字标签：浅灰底 + 深灰字（对应 HTML .sr-tag--light）。
        作为单独段落渲染，前后留空格让标签块视觉上独立。"""
        text = node.get_text(" ", strip=True)
        if not text:
            return
        p = container.add_paragraph()
        run = p.add_run(f" {text} ")
        run.bold = False
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("1A1F36")
        self._set_run_shading(run, "EDF1F7")

    def _handle_muted_note(self, node, container):
        """图表注释/小字灰体段（对应 HTML .report-chart-note.sr-p.muted）：
        小字 + 浅灰色字 + 不加粗。"""
        text = node.get_text(" ", strip=True)
        if not text:
            return
        p = container.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string("8A8F9A")

    def _handle_sub_paragraph(self, node, container):
        """子段落（对应 HTML .report-body.sr-p.sr-p--sub）：
        整段浅灰底 + 段前/段后留白，视觉与正文段落做区分。"""
        p = container.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        self._render_inline_with_br(p, node)
        if not p.runs:
            p._element.getparent().remove(p._element)
            return
        for run in p.runs:
            self._set_run_shading(run, "F5F7FA")

    def _handle_risk_harm_box(self, node, container):
        """风险危害解读框（对应 HTML .sr-risk-harm-box，Figma 425:11151）："""
        # 颜色常量（与 HTML CSS 变量保持一致）
        box_bg = "EDF1F7"      # --color-graphite-l40 = --sr-surface-muted
        badge_bg = "1C6EFF"    # sr-risk-harm-box__badge 渐变蓝色（简化为纯色）
        badge_fg = "FFFFFF"     # badge 白字
        desc_color = "6B7A99"  # --tx2 = --sr-text-secondary

        # Word 两个表格之间强制要有段落。前一个兄弟节点是 table（C2/病毒/漏洞
        # 利用事件表）时，Word 会把两个相邻表格自动合并成一个，导致 risk-harm-box
        # 被吸进事件表里。这里无条件插入一个分隔段落，提供 6pt 间隙，并避免合并。
        sep_p = container.add_paragraph()
        sep_p.paragraph_format.space_before = Pt(0)
        sep_p.paragraph_format.space_after = Pt(6)
        sep_p.paragraph_format.line_spacing = 1
        sep_run = sep_p.add_run("")
        sep_run.font.size = Pt(1)

        # 创建 1×1 表格作为外框
        table = container.add_table(rows=1, cols=1)
        table.autofit = False
        table.allow_autofit = False
        # 表格宽度 = 页面内容宽度（与前面事件表对齐左右边）
        cfg = self.config.get("docx", {})
        page_w_mm = cfg.get("page_width_mm", 210)
        margin_mm = cfg.get("margin_mm", 20)
        content_w_mm = page_w_mm - 2 * margin_mm
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        # 移除 python-docx 默认套用的 tblStyle（带灰色边框），改由我们直接控制
        for style_tag in ("w:tblStyle",):
            existing = tblPr.findall(qn(style_tag))
            for e in existing:
                tblPr.remove(e)
        # 移除默认 tblBorders（python-docx 默认带 single D3D7DE 边框）
        existing_borders = tblPr.findall(qn("w:tblBorders"))
        for e in existing_borders:
            tblPr.remove(e)
        # tblInd=0（左对齐页面内容区左边，与前面事件表对齐）
        existing_ind = tblPr.findall(qn("w:tblInd"))
        for e in existing_ind:
            tblPr.remove(e)
        tblInd = OxmlElement("w:tblInd")
        tblInd.set(qn("w:w"), "0")
        tblInd.set(qn("w:type"), "dxa")
        tblPr.append(tblInd)
        # tblLayout=fixed
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        tblPr.append(layout)
        # tblW 总宽
        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:type"), "dxa")
        tblW.set(qn("w:w"), str(int(round(content_w_mm * 1440 / 25.4))))
        tblPr.append(tblW)
        # 单元格 margins（模拟 box padding：上下 8pt 左右 12pt）
        tblCellMar = OxmlElement("w:tblCellMar")
        # 左右 padding 用 108 twips（与事件表一致，确保 risk-harm-box 内容区左右边
        # 与事件表对齐）；上下 padding 160 twips 保持框内留白
        for side, twips in (("top", 160), ("bottom", 160), ("left", 108), ("right", 108)):
            mar = OxmlElement(f"w:{side}")
            mar.set(qn("w:w"), str(twips))
            mar.set(qn("w:type"), "dxa")
            tblCellMar.append(mar)
        tblPr.append(tblCellMar)
        # 表格边框：设为 nil（无边框），HTML .sr-risk-harm-box 视觉只有底色无外框线
        tblBorders = OxmlElement("w:tblBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "nil")
            tblBorders.append(border)
        tblPr.append(tblBorders)
        # tblGrid + gridCol
        tblGrid = OxmlElement("w:tblGrid")
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(int(round(content_w_mm * 1440 / 25.4))))
        tblGrid.append(gridCol)
        tbl.insert(1, tblGrid)
        cell = table.rows[0].cells[0]
        # 单元格底色
        self._set_cell_shading(cell, box_bg)

        # 复用单元格默认首段（python-docx 创建单元格时会自带一个空段落），
        # 避免再 add_paragraph 造成首段空行 + badge 段两段结构。
        content_p = cell.paragraphs[0]
        content_p.paragraph_format.space_before = Pt(0)
        content_p.paragraph_format.space_after = Pt(0)
        content_p.paragraph_format.line_spacing = 1.15
        badge_text = " 风险危害 "  # 左右各加一个空格模拟 padding
        badge_run = content_p.add_run(badge_text)
        badge_run.bold = True
        badge_run.font.size = Pt(10.5)
        badge_run.font.color.rgb = RGBColor.from_string(badge_fg)
        self._set_run_shading(badge_run, badge_bg)

        # 描述文本：取 node 内 <p> 元素文本，作为同段后续 run
        desc_text = ""
        p_tag = node.find("p")
        if p_tag is not None:
            desc_text = p_tag.get_text(" ", strip=True)
        else:
            # 兜底：取 badge span 之外的所有文本
            desc_text = node.get_text(" ", strip=True).replace("风险危害", "", 1).strip()
        if desc_text:
            # badge 和描述之间加一个空格分隔
            sep_run = content_p.add_run(" ")
            sep_run.font.size = Pt(10.5)
            desc_run = content_p.add_run(desc_text)
            desc_run.font.size = Pt(10.5)
            desc_run.font.color.rgb = RGBColor.from_string(desc_color)


    def _handle_risk_card_head(self, node, container, reuse_first_para=False):
        """风险卡头部：把 sr-tag 和 sr-risk-card__title 拼到同一段落，
        单行带底色块呈现（对应 HTML header.sr-risk-card__head）。

        reuse_first_para=True 时复用 cell 默认的第一个空段落，避免出现额外
        空段（用户 mark"这个换行不要"）。

        标题字号 13pt（>11pt，满足用户 mark"这行字体字号要大于11"）。

        mark 调整（2026-08）：
          - 标签（威胁运营/威胁预防/防护有效性）保持可编辑文本（run 底纹），
            不做圆角图片。
          - 标签不再顶格：头段落设 space_before 上间距（Word 单元格默认
            无内边距，用段前间距模拟 HTML 的 padding-top）。"""
        if reuse_first_para and container.is_cell and container.cell.paragraphs:
            p = container.cell.paragraphs[0]
            # 清掉默认空 run
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
        else:
            p = container.add_paragraph()
        # 头段落：上间距（非顶格）
        p.paragraph_format.space_before = Pt(_RISK_CARD_HEAD_TAG_SPACE_BEFORE)
        for child in node.children:
            if isinstance(child, NavigableString):
                text = str(child).strip()
                if text:
                    p.add_run(text + " ")
            elif isinstance(child, Tag):
                classes = child.get("class") or []
                if child.name in ("h6", "h5", "h4", "h3", "h2", "h1") or "sr-risk-card__title" in classes:
                    title_text = child.get_text(" ", strip=True)
                    if title_text:
                        run = p.add_run(title_text)
                        run.bold = True
                        run.font.size = Pt(13)  # >11pt（用户 mark）
                        run.font.color.rgb = RGBColor.from_string("1A1F36")
                elif "sr-tag--light" in classes:
                    # sr-tag--light 浅底深字样式；底色按次要类（--medium/--blue/--success 等）
                    # 单独分配（用户要求"按照标签单独分配一些合适的底色"）
                    # bg 用浅色调（对应 HTML 中 .sr-tag--light + .sr-tag--X 的语义）
                    light_bg_map = {
                        "sr-tag--critical":    ("FFD3CC", "CF171D"),  # 浅红 + 深红字
                        "sr-tag--high":        ("FFE0BF", "FA721B"),  # 浅橙 + 深橙字
                        "sr-tag--medium":      ("FFE9CC", "D6860D"),  # 浅黄 + 深黄字
                        "sr-tag--medium-low":  ("FFF0BF", "D6860D"),  # 浅米黄 + 深黄字
                        "sr-tag--low":         ("E8F4FF", "1C6EFF"),  # 浅蓝 + 蓝字
                        "sr-tag--info":        ("E8F4FF", "1C6EFF"),  # 浅蓝 + 蓝字
                        "sr-tag--success":     ("CCFFE7", "12A679"),  # 浅绿 + 深绿字
                        "sr-tag--blue":        ("E8F4FF", "1C6EFF"),  # 浅蓝 + 蓝字
                    }
                    bg_hex, fg_hex = ("EDF1F7", "1A1F36")  # 默认浅灰
                    for c in classes:
                        if c in light_bg_map:
                            bg_hex, fg_hex = light_bg_map[c]
                            break
                    text = child.get_text(" ", strip=True)
                    if text:
                        run = p.add_run(f" {text} ")
                        run.font.size = Pt(10.5)  # 五号（用户 mark）
                        run.font.bold = True
                        run.font.color.rgb = RGBColor.from_string(fg_hex)
                        self._set_run_shading(run, bg_hex)
                elif any(c in ("sr-tag--success", "sr-tag--info", "sr-tag--blue",
                               "sr-tag--critical", "sr-tag--high", "sr-tag--medium",
                               "sr-tag--medium-low", "sr-tag--low") for c in classes):
                    # 用对应 tag handler 的颜色，作为 inline run 写入
                    tag_key = next((f"tag_{c.replace('sr-tag--', '')}" for c in classes
                                    if c.startswith("sr-tag--") and c != "sr-tag--light"), None)
                    text = child.get_text(" ", strip=True)
                    if text and tag_key:
                        bg, fg = _TAG_COLORS.get(tag_key, ("6F7785", "FFFFFF"))
                        run = p.add_run(f" {text} ")
                        run.font.size = Pt(10.5)  # 五号（与浅底标签一致）
                        run.bold = True
                        run.font.color.rgb = RGBColor.from_string(fg)
                        self._set_run_shading(run, bg)
                else:
                    text = child.get_text(" ", strip=True)
                    if text:
                        p.add_run(text + " ")

    def _handle_top5_asset_ip_row(self, node, container):
        """Top5 资产 IP 行：把 IP 和"未托管/已托管"标签拼到同一段落，
        标签渲染为灰色小字 + 括号包含（对应 HTML .sr-top5-asset-ip-row）。
        IP 不加粗（对应 HTML .sr-top5-asset-ip 的 font-weight:400）。
        例：<div class="sr-top5-asset-ip-row"><span class="sr-top5-asset-ip">10.1.2.3</span><span class="sr-tag sr-tag--light sr-tag--medium">未托管</span></div>
        → Word: "10.1.2.3 (未托管)" 其中 "(未托管)" 是灰色 8.5pt 小字。"""
        p = container.add_paragraph()
        for child in node.children:
            if isinstance(child, NavigableString):
                text = str(child).strip()
                if text:
                    p.add_run(text + " ")
            elif isinstance(child, Tag):
                classes = child.get("class") or []
                if "sr-top5-asset-ip" in classes:
                    ip_text = child.get_text(" ", strip=True)
                    if ip_text:
                        ip_run = p.add_run(ip_text)
                        ip_run.bold = False
                elif any(c.startswith("sr-tag") for c in classes):
                    tag_text = child.get_text(" ", strip=True)
                    if tag_text:
                        # 灰色小字 + 括号
                        tag_run = p.add_run(f" ({tag_text})")
                        tag_run.font.size = Pt(8.5)
                        tag_run.font.color.rgb = RGBColor.from_string("8A8F9A")
                else:
                    text = child.get_text(" ", strip=True)
                    if text:
                        p.add_run(text + " ")

    def _handle_grade_badge(self, node, container):
        """评级徽章：居中加粗。"""
        text = node.get_text(" ", strip=True)
        if not text:
            return
        p = container.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16)

    def _map_tag(self, node, container, tag_key):
        """风险标签：红/橙/黄底色 + 白字 run。"""
        text = node.get_text(" ", strip=True)
        if not text:
            return
        bg, fg = _TAG_COLORS.get(tag_key, ("6F7785", "FFFFFF"))
        p = container.add_paragraph()
        run = p.add_run(f" {text} ")
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(fg)
        self._set_run_shading(run, bg)

    def _map_heading(self, node, container, level=1):
        text = node.get_text(" ", strip=True)
        if not text:
            return
        p = container.add_heading(text, level=level)
        # 兜底：清掉该段落的 numPr，避免 Word 自动加序号叠加成 "1. 1. xxx"
        self._strip_paragraph_num_pr(p)

    def _map_paragraph(self, node, container):
        # 段落内部可能包含 <br> / 块级标签 / 行内标签，需要按 <br> 拆段，
        # 与表格单元格 _render_cell_text_with_br 保持一致。
        p = container.add_paragraph()
        self._render_inline_with_br(p, node)
        # 解析 inline style 中的 text-indent:Nem，设 Word 首行缩进
        self._apply_text_indent(node, p, container)
        # 风险卡内的 p 正文首行缩进 2em（用户 mark"首行缩进"）
        if getattr(self, "_in_risk_card", False) and not p.paragraph_format.first_line_indent:
            # 检查父节点是否是 sr-risk-field（"问题描述"/"风险影响"等下面的正文）
            parent = node.parent
            if parent is not None and "sr-risk-field" in (parent.get("class") or []):
                from docx.shared import Mm as _Mm
                # 2em ≈ 21pt ≈ 7.4mm，按 10.5pt 字号近似
                p.paragraph_format.first_line_indent = _Mm(7.4)
        if not p.runs:
            # 没有任何 run 说明未写入内容，移除该空段
            p._element.getparent().remove(p._element)
        else:
            # 普通段落也可能继承到 List Number 样式导致自动编号，兜底清掉
            self._strip_paragraph_num_pr(p)

    def _apply_text_indent(self, node, paragraph, container):
        """从 node 的 style 中解析 text-indent:Nem 和 margin-top:Npx，设 Word 首行缩进和段前间距。"""
        style = node.get("style") if isinstance(node, Tag) else None
        if not style:
            return
        # 解析 margin-top 作为 Word space_before（1px ≈ 0.75pt）
        m = re.search(r'margin-top\s*:\s*([\d.]+)px', style)
        if m:
            px = float(m.group(1))
            paragraph.paragraph_format.space_before = Pt(round(px * 0.75, 1))
        m = re.search(r'text-indent\s*:\s*([\d.]+)\s*em', style)
        if not m:
            return
        em = float(m.group(1))
        # 取当前正文字号作为 1em 的磅值
        size_pt = 10.5
        fonts_cfg = getattr(container, "fonts", {}) or {}
        normal_cfg = fonts_cfg.get("normal") if isinstance(fonts_cfg, dict) else None
        if isinstance(normal_cfg, dict) and normal_cfg.get("size_pt"):
            try:
                size_pt = float(normal_cfg["size_pt"])
            except (TypeError, ValueError):
                pass
        paragraph.paragraph_format.first_line_indent = Pt(em * size_pt)

    def _apply_run_fmt(self, run, fmt):
        """根据 fmt 描述给 run 应用 bold/italic/color。
        fmt 可为 None、字符串 "bold"/"italic" 或 dict {"bold","italic","color"}。"""
        if fmt is None:
            return
        if isinstance(fmt, str):
            if fmt == "bold":
                run.bold = True
            elif fmt == "italic":
                run.italic = True
            return
        if isinstance(fmt, dict):
            if fmt.get("bold"):
                run.bold = True
            if fmt.get("italic"):
                run.italic = True
            color = fmt.get("color")
            if color:
                try:
                    run.font.color.rgb = RGBColor.from_string(color)
                except Exception:
                    pass
            size = fmt.get("size")
            if size is not None:
                try:
                    run.font.size = Pt(size)
                except Exception:
                    pass

    def _render_grade_badge_inline(self, paragraph, node):
        """将 <span class="sr-grade sr-grade--良">良</span> 作为 inline run 渲染：
        浅底色 + 深色字 + 9pt 加粗，左右留空格模拟徽章 padding。"""
        text = node.get_text(" ", strip=True)
        if not text:
            return
        classes = node.get("class", []) or []
        if isinstance(classes, str):
            classes = classes.split()
        grade_key = next((c for c in classes if c.startswith("sr-grade--")), None)
        bg, fg = _GRADE_COLORS.get(grade_key, _GRADE_DEFAULT)
        # 徽章前后留一个空格模拟 HTML 的 padding:2px 10px
        run = paragraph.add_run(f" {text} ")
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor.from_string(fg)
        self._set_run_shading(run, bg)

    def _render_inline_with_br(self, paragraph, node):
        """把 node 子节点按 <br> / 块级标签拆分写入段落，<br> 映射为软换行。
        <strong>/<b> → 加粗；<em>/<i> → 斜体；语义保留到独立 run。"""
        block_tags = {"div", "p"}
        bold_tags = {"strong", "b"}
        italic_tags = {"em", "i"}
        parts = []
        current_text = []
        children = list(node.children)
        for i, child in enumerate(children):
            if isinstance(child, Tag) and child.name == "br":
                parts.append(("text", "".join(current_text), None))
                parts.append(("br", None, None))
                current_text = []
            elif isinstance(child, Tag) and child.name in block_tags:
                parts.append(("text", "".join(current_text), None))
                parts.append(("text", child.get_text(" ", strip=True), None))
                if i < len(children) - 1:
                    parts.append(("br", None, None))
                current_text = []
            elif isinstance(child, NavigableString):
                current_text.append(str(child))
            elif isinstance(child, Tag) and child.name in bold_tags:
                if current_text:
                    parts.append(("text", "".join(current_text), None))
                    current_text = []
                parts.append(("text", child.get_text(" ", strip=True), "bold"))
            elif isinstance(child, Tag) and child.name in italic_tags:
                if current_text:
                    parts.append(("text", "".join(current_text), None))
                    current_text = []
                parts.append(("text", child.get_text(" ", strip=True), "italic"))
            elif isinstance(child, Tag):
                # 评级徽章 sr-grade：作为带底色 inline run 渲染到同一段落
                child_classes = child.get("class", []) or []
                if isinstance(child_classes, str):
                    child_classes = child_classes.split()
                if any(c == "sr-grade" or c.startswith("sr-grade--") for c in child_classes):
                    if current_text:
                        parts.append(("text", "".join(current_text), None))
                        current_text = []
                    parts.append(("grade_badge", child, None))
                else:
                    # 其他 span/容器标签：递归提取，保留嵌套 strong/b/em/i 的 bold/italic 语义
                    self._collect_inline_parts(child, current_text, parts, None)
        if current_text:
            parts.append(("text", "".join(current_text), None))
        for ptype, content, fmt in parts:
            if ptype == "text":
                text = (content or "").strip()
                if text:
                    run = paragraph.add_run(text)
                    self._apply_run_fmt(run, fmt)
            elif ptype == "br":
                run = paragraph.add_run()
                run.add_break()
            elif ptype == "grade_badge":
                self._render_grade_badge_inline(paragraph, content)

    def _collect_inline_parts(self, node, current_text, parts, inherited_fmt):
        """递归收集 node 子节点的 inline 文本，保留嵌套的 strong/b/em/i 与 color 语义。

        - 遇到 NavigableString 追加到 current_text
        - 遇到 strong/b：先 flush 当前累计文本，再开新 bold 文本段
        - 遇到 em/i：同上但 fmt=italic
        - 遇到 sr-text-danger span：以 {color: "CF171D"} 作为 inherited_fmt 递归子节点
        - 遇到其他 Tag：递归子节点（fmt 不变）

        inherited_fmt 可为 None 或 dict {"bold": bool, "italic": bool, "color": str}。
        非 None 时，bold/italic 会作用于该子树下的普通文本与 strong/em 文本，
        color 也会传播到这些 run。"""
        bold_tags = {"strong", "b"}
        italic_tags = {"em", "i"}
        # 若 node 自身就是 sr-text-danger 容器，先把已累计文本用外层 fmt flush，
        # 再以新 color 注入 inherited_fmt，避免外层文本被错误染红
        if isinstance(node, Tag):
            node_classes = node.get("class", []) or []
            if isinstance(node_classes, str):
                node_classes = node_classes.split()
            if "sr-text-danger" in node_classes:
                if current_text:
                    parts.append(("text", "".join(current_text), inherited_fmt))
                    current_text.clear()
                new_fmt = {"color": "CF171D"}
                if isinstance(inherited_fmt, dict):
                    if inherited_fmt.get("bold"):
                        new_fmt["bold"] = True
                    if inherited_fmt.get("italic"):
                        new_fmt["italic"] = True
                inherited_fmt = new_fmt
        for child in node.children:
            if isinstance(child, Tag) and child.name == "br":
                # 软换行：flush 已累计文本，追加 br 标记（写入段落后映射为 <w:br/>）
                if current_text:
                    parts.append(("text", "".join(current_text), inherited_fmt))
                    current_text.clear()
                parts.append(("br", None, None))
            elif isinstance(child, NavigableString):
                current_text.append(str(child))
            elif isinstance(child, Tag) and child.name in bold_tags:
                if current_text:
                    parts.append(("text", "".join(current_text), inherited_fmt))
                    current_text.clear()
                fmt = {"bold": True}
                if isinstance(inherited_fmt, dict):
                    if inherited_fmt.get("italic"):
                        fmt["italic"] = True
                    if inherited_fmt.get("color"):
                        fmt["color"] = inherited_fmt["color"]
                parts.append(("text", child.get_text(" ", strip=True), fmt))
            elif isinstance(child, Tag) and child.name in italic_tags:
                if current_text:
                    parts.append(("text", "".join(current_text), inherited_fmt))
                    current_text.clear()
                fmt = {"italic": True}
                if isinstance(inherited_fmt, dict):
                    if inherited_fmt.get("bold"):
                        fmt["bold"] = True
                    if inherited_fmt.get("color"):
                        fmt["color"] = inherited_fmt["color"]
                parts.append(("text", child.get_text(" ", strip=True), fmt))
            elif isinstance(child, Tag):
                child_classes = child.get("class", []) or []
                if isinstance(child_classes, str):
                    child_classes = child_classes.split()
                if "sr-text-danger" in child_classes:
                    # 红色字体容器：先把已累计文本用外层 fmt flush，再以新 color 递归子节点
                    if current_text:
                        parts.append(("text", "".join(current_text), inherited_fmt))
                        current_text.clear()
                    new_fmt = {"color": "CF171D"}
                    if isinstance(inherited_fmt, dict):
                        if inherited_fmt.get("bold"):
                            new_fmt["bold"] = True
                        if inherited_fmt.get("italic"):
                            new_fmt["italic"] = True
                    self._collect_inline_parts(child, current_text, parts, new_fmt)
                else:
                    self._collect_inline_parts(child, current_text, parts, inherited_fmt)

    def _map_list(self, node, container, ordered=False):
        from docx.shared import Mm as _Mm
        for li in node.find_all("li", recursive=False):
            # 若 <li> 自带手写序号（如 sr-risk-step__num / sr-risk-step），
            # 跳过 Word 自动编号样式。把"手写序号 + 标题"作为单独段落，再
            # 递归处理内层嵌套 <ul>/<ol>，避免 get_text 把整棵子树拍平成一段。
            li_classes = li.get("class") or []
            has_handwritten_num = (
                "sr-risk-step" in li_classes
                or li.find(class_="sr-risk-step__num") is not None
            )
            if has_handwritten_num:
                self._map_risk_step(li, container)
                continue
            text = li.get_text(" ", strip=True)
            if not text:
                continue
            style = "List Number" if ordered else "List Bullet"
            p = container.add_paragraph(text, style=style)
            # 用户 mark"所有项目符号这里，整体增加一个缩进"：
            # 风险卡步骤下的子列表项加 10mm 左缩进（5mm step + 5mm 子列表）
            if getattr(self, "_in_risk_step_list", False):
                p.paragraph_format.left_indent = _Mm(10)

    def _map_risk_step(self, li, container):
        """渲染 sr-risk-step：手写序号 + 标题段 + 嵌套列表段落。

        HTML 结构：
          <li class="sr-risk-step">
            <span class="sr-risk-step__num">1</span>
            <div class="sr-risk-step__content">
              <p class="sr-risk-step__title">网络防护</p>
              <ul class="sr-risk-step__list">
                <li>开启防火墙云情报网关订阅</li>
                ...
              </ul>
            </div>
          </li>
        生成：
          - 加粗段落"1 网络防护"（无 Word 自动序号，左缩进 5mm）
          - 嵌套 <ul> 各 <li> 作为 List Bullet 段落（整体再加 5mm 缩进 = 10mm）
        """
        from docx.shared import Mm as _Mm
        num_span = li.find(class_="sr-risk-step__num")
        num_text = num_span.get_text(strip=True) if num_span else ""
        # 找标题
        title_tag = li.find(class_="sr-risk-step__title")
        title_text = title_tag.get_text(" ", strip=True) if title_tag else ""
        # 生成"序号 + 标题"加粗段
        head_text = f"{num_text} {title_text}".strip() if title_text else num_text
        if head_text:
            p = container.add_paragraph()
            # 用户 mark"所有编号这里，整体增加一个缩进"：左缩进 5mm
            p.paragraph_format.left_indent = _Mm(5)
            run = p.add_run(head_text)
            run.bold = True
            self._strip_paragraph_num_pr(p)
        # 处理内层嵌套列表（sr-risk-step__list 下的 <ul>/<ol>）
        nested_lists = li.find_all(["ul", "ol"], recursive=True)
        # 排除被嵌套 sr-risk-step（不会有，因为 sr-risk-step 在 li 上）
        # 标记：风险卡步骤下的子列表需要额外加 5mm 缩进
        self._in_risk_step_list = True
        for nested in nested_lists:
            # 避免重复处理：只处理直接属于本 li 的、不被更深层 sr-risk-step 包裹的
            # 简化：用 recursive=True 找到所有，但跳过 nested 的子 ul/ol
            if nested.find_parent(class_="sr-risk-step") is not li:
                continue
            self._map_list(nested, container, ordered=(nested.name == "ol"))
        self._in_risk_step_list = False

    def _image_max_size(self):
        """计算页面可用宽高（mm），用于限制图片不超页。"""
        cfg = self.config["docx"]
        content_w = cfg.get("page_width_mm", 297) - 2 * cfg.get("margin_mm", 20)
        content_h = cfg.get("page_height_mm", 210) - 2 * cfg.get("margin_mm", 20)
        return content_w, content_h

    def _fit_image_size(self, img_source):
        """根据图片像素尺寸与 scale，计算 Mm(width) 与 Mm(height)，并按页面
        可用宽高做等比缩放，保证既不超宽也不超高。

        img_source 可为文件路径字符串或 BytesIO 流。返回 (width_emu, height_emu)
        或 (None, None) 表示无法读取尺寸时退回 None。
        """
        try:
            with Image.open(img_source) as im:
                px_w, px_h = im.size
        except Exception as e:
            _log(f"读取图片尺寸失败，按默认宽度处理: {e}", "WARNING")
            return None, None
        # PNG 截图 scale=2，实际显示 DPI 默认按 96 处理：1px ≈ 0.2645 mm
        # 用 px → mm 直接换算并让 python-docx 用 Mm 落地
        mm_w = px_w * 25.4 / 96.0
        mm_h = px_h * 25.4 / 96.0
        max_w_mm, max_h_mm = self._image_max_size()
        # 优先约束宽度，再判断高度是否超页
        width_mm = min(mm_w, max_w_mm)
        scale = width_mm / mm_w if mm_w > 0 else 1.0
        height_mm = mm_h * scale
        # 若高度仍超页，按高度再缩一次
        if height_mm > max_h_mm:
            height_mm = max_h_mm
            scale = height_mm / mm_h if mm_h > 0 else 1.0
            width_mm = mm_w * scale
        return Mm(width_mm), Mm(height_mm)

    def _map_image(self, node, container):
        """处理 base64 data:image 与截图占位。

        截图占位 <img data-snapshot="...">：若 _snapshot_titles[snap] 存在，
        先在 Word 里写小标题段（与 HTML 中 .chart-title 同样的样式），
        再插图，最后在图下加居中"图 N-M 标题"形式的命名段。
        """
        snap = node.get("data-snapshot")
        if snap and snap in self._snapshot_map:
            img_path = self._snapshot_map[snap]
            if not os.path.exists(img_path):
                _log(f"截图文件不存在（忽略）: {img_path}", "WARNING")
                return
            titles = getattr(self, "_snapshot_titles", {}) or {}
            chart_title = (titles.get(snap) or "").strip()
            chart_ids = getattr(self, "_snapshot_chart_ids", {}) or {}
            chart_id = (chart_ids.get(snap) or "").strip()
            # 这 2 个图缩小 50% 居中（用户要求）
            # 注：3.2.3 的 m2-ring2 / m2-bar-sys 已合成一张图整体展示，不再独立出现
            shrink_ids = {"m3-web-top5-bar", "m3-nonweb-top5-bar"}
            should_shrink = chart_id in shrink_ids
            # 标题保留在图里的 snap：不写图前小标题（标题已在图里），但仍写图后 caption
            keep_title_snaps = getattr(self, "_snapshot_keep_title", set()) or set()
            skip_inline_title = snap in keep_title_snaps
            # grid-2 合成图的特殊图注（如 3.3.1 Web/非Web top5 合成图）
            grid_captions = getattr(self, "_snapshot_grid_captions", {}) or {}
            grid_caption = (grid_captions.get(snap) or "").strip()
            # 图前：写小标题（沿用 HTML .chart-title 视觉：左对齐、加粗、11pt）
            # 跳过条件：标题保留在图里 / grid-2 合成图（标题已在图里）
            if chart_title and not skip_inline_title and not grid_caption:
                title_p = container.add_paragraph()
                title_p.paragraph_format.space_before = Pt(4)
                title_p.paragraph_format.space_after = Pt(2)
                t_run = title_p.add_run(chart_title)
                t_run.bold = True
                t_run.font.size = Pt(11)
                t_run.font.color.rgb = RGBColor.from_string("1A1F36")
            w, h = self._fit_image_size(img_path)
            if should_shrink:
                # 缩小 50%：宽度减半
                if w is None or h is None:
                    content_w = self.config["docx"]["page_width_mm"] - 2 * self.config["docx"]["margin_mm"]
                    pic_w = Mm(content_w / 2)
                    pic_h = None
                else:
                    pic_w = Mm(int(w.mm / 2))
                    pic_h = Mm(int(h.mm / 2))
                pic_p = container.add_picture(img_path, width=pic_w, height=pic_h)
                # 居中
                pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif w is None or h is None:
                content_w = self.config["docx"]["page_width_mm"] - 2 * self.config["docx"]["margin_mm"]
                pic_p = container.add_picture(img_path, width=Mm(content_w))
            else:
                pic_p = container.add_picture(img_path, width=w, height=h)
            # 图片段段后间距强制为 0，让下方图注紧贴图片
            if pic_p is not None:
                pic_p.paragraph_format.space_before = Pt(0)
                pic_p.paragraph_format.space_after = Pt(0)
            # 图后：居中"图 N-M 标题"形式的命名段（编号按图出现顺序自增）
            # grid-2 合成图：用 grid_caption 作为图注文字（无图前小标题）
            # no_caption 集：完全跳过图后 caption（标题已在图里，且不希望加图注）
            no_caption_snaps = getattr(self, "_snapshot_no_caption", set()) or set()
            skip_caption = snap in no_caption_snaps
            if grid_caption and not skip_caption:
                self._chart_caption_index = getattr(self, "_chart_caption_index", 0) + 1
                caption_p = container.add_paragraph()
                caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_p.paragraph_format.space_before = Pt(2)
                caption_p.paragraph_format.space_after = Pt(0)
                c_run = caption_p.add_run(f"图 {self._chart_caption_index} {grid_caption}")
                c_run.font.size = Pt(9)
                c_run.font.color.rgb = RGBColor.from_string("6F7785")
            elif chart_title and not skip_caption:
                self._chart_caption_index = getattr(self, "_chart_caption_index", 0) + 1
                caption_p = container.add_paragraph()
                caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_p.paragraph_format.space_before = Pt(2)
                caption_p.paragraph_format.space_after = Pt(0)
                c_run = caption_p.add_run(f"图 {self._chart_caption_index} {chart_title}")
                c_run.font.size = Pt(9)
                c_run.font.color.rgb = RGBColor.from_string("6F7785")
            return
        src = node.get("src") or ""
        if src.startswith("data:image/"):
            head, b64 = src.split(",", 1) if "," in src else ("", src)
            try:
                raw = base64.b64decode(b64)
                stream = io.BytesIO(raw)
                w, h = self._fit_image_size(stream)
                stream.seek(0)
                if w is None or h is None:
                    container.add_picture(stream, width=Mm(80))
                else:
                    container.add_picture(stream, width=w, height=h)
            except Exception as e:
                _log(f"base64 图片解码失败（忽略）: {e}", "WARNING")
            return
        if src:
            try:
                if src.startswith("file://"):
                    img_path = src[7:].lstrip("/")
                else:
                    img_path = str((self.input_path.parent / src).resolve())
                if os.path.exists(img_path):
                    w, h = self._fit_image_size(img_path)
                    if w is None or h is None:
                        container.add_picture(img_path, width=Mm(80))
                    else:
                        container.add_picture(img_path, width=w, height=h)
            except Exception as e:
                _log(f"外部图片加载失败（忽略）: {src} ({e})", "WARNING")

    def _map_table(self, node, container):
        """映射 HTML table 到 docx table，处理 colspan/rowspan 以及 <br> 换行。"""
        rows = node.find_all("tr", recursive=True)
        if not rows:
            return
        # 文档信息表（sr-copyright-meta）单独设垂直居中，其他表保持默认顶部对齐
        node_classes = node.get("class") or []
        is_doc_info_table = "sr-copyright-meta" in node_classes
        max_cols = 0
        grid = []
        has_any_th = False
        for tr in rows:
            cells = tr.find_all(["td", "th"], recursive=False)
            col_sum = 0
            row_data = []
            for c in cells:
                try:
                    colspan = int(c.get("colspan", 1) or 1)
                except (ValueError, TypeError):
                    colspan = 1
                try:
                    rowspan = int(c.get("rowspan", 1) or 1)
                except (ValueError, TypeError):
                    rowspan = 1
                if c.name == "th":
                    has_any_th = True
                row_data.append((c, colspan, rowspan))
                col_sum += colspan
            max_cols = max(max_cols, col_sum)
            grid.append(row_data)
        if max_cols == 0:
            return
        table = container.add_table(rows=len(grid), cols=max_cols)
        occupied = {}
        for ri, row_data in enumerate(grid):
            ci = 0
            is_header_row = (ri == 0 and not has_any_th) or all(
                soup.name == "th" for soup, _, _ in row_data
            )
            for cell_soup, colspan, rowspan in row_data:
                while (ri, ci) in occupied:
                    ci += 1
                cell = table.cell(ri, ci)
                self._render_cell_text_with_br(cell, cell_soup)
                # 所有表格单元格默认垂直居中（水平对齐保持原状）
                try:
                    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                except Exception as e:
                    _log(f"设置单元格垂直居中失败: {e}", "WARNING")
                if colspan > 1:
                    for k in range(1, colspan):
                        try:
                            merged = cell.merge(table.cell(ri, ci + k))
                            cell = merged
                        except Exception:
                            pass
                if rowspan > 1:
                    for r in range(1, rowspan):
                        for k in range(colspan):
                            occupied[(ri + r, ci + k)] = True
                if is_header_row or cell_soup.name == "th":
                    self._apply_table_header_cell_style(cell)
                    # 表头默认不换行（用户要求：表头不换行）
                    self._set_cell_no_wrap(cell)
                # 单元格 class 含 sr-no-wrap 时不换行（用于时间列等需要保持完整的字段）
                if _has_class(cell_soup, "sr-no-wrap"):
                    self._set_cell_no_wrap(cell)
                ci += colspan
        self._set_table_column_widths_by_grid(table, grid, node)
        self._set_table_borders(table)
        # 附录（#sec-appendix）内的表格：表头行（全 th）水平居中显示。
        # HTML 中 .sr-tbl th 默认 text-align:left，Word 单元格段落未设对齐时
        # 同样继承 left；这里把附录内表格的表头单元格段落统一居中。
        if node.find_parent(id="sec-appendix") is not None:
            for row_index, row_data in enumerate(grid):
                if row_data and all(cell_soup is not None and cell_soup.name == "th"
                                    for cell_soup, _, _ in row_data):
                    for cell in table.rows[row_index].cells:
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 仅局部：sr-top5-tbl（2.2节 TOP5风险资产表）左对齐 + 2mm缩进，与HTML对齐一致
        if "sr-top5-tbl" in node_classes:
            try:
                tbl = table._tbl
                tblPr = tbl.find(qn('w:tblPr'))
                if tblPr is None:
                    tblPr = OxmlElement('w:tblPr')
                    tbl.insert(0, tblPr)
                for tag in ('w:jc', 'w:tblInd'):
                    old = tblPr.find(qn(tag))
                    if old is not None:
                        tblPr.remove(old)
                jc = OxmlElement('w:jc')
                jc.set(qn('w:val'), 'left')
                tblPr.append(jc)
                tblInd = OxmlElement('w:tblInd')
                tblInd.set(qn('w:w'), str(int(Mm(2).twips)))
                tblInd.set(qn('w:type'), 'dxa')
                tblPr.append(tblInd)
            except Exception:
                pass

        # 组件策略检查异常项表（sr-component-check-tbl）：第 1 列（序号）与
        # 最后一列（风险说明）在自动分配的列宽基础上各加宽一个汉字宽度
        # （10.5pt 中文 ≈ 3.7mm），让这两列能多容纳一个汉字；中间列等比收缩，
        # 总宽保持页面内容宽度，避免 Word 按内容撑开溢出。
        if "sr-component-check-tbl" in node_classes:
            try:
                cfg_w = self.config.get("docx", {})
                content_w_mm = cfg_w.get("page_width_mm", 210) - 2 * cfg_w.get("margin_mm", 20)
                tbl = table._tbl
                tblGrid = tbl.find(qn("w:tblGrid"))
                if tblGrid is None:
                    raise RuntimeError("no tblGrid")
                grid_cols = tblGrid.findall(qn("w:gridCol"))
                ncols = len(grid_cols)
                if ncols < 3:
                    raise RuntimeError("too few cols")
                # 自动分配后的各列宽度（twips -> mm）
                cur = [int(c.get(qn("w:w"), 0)) * 25.4 / 1440.0 for c in grid_cols]
                char_mm = 3.7  # 10.5pt 一个汉字 ≈ 3.7mm
                # 头尾列各加一个汉字宽
                cur[0] += char_mm
                cur[-1] += char_mm
                total = sum(cur)
                if total > content_w_mm and total > 0:
                    # 先收缩中间列，尽量保住头尾的加宽
                    overflow = total - content_w_mm
                    middle_old = sum(cur[1:-1])
                    if middle_old > overflow:
                        scale = (middle_old - overflow) / middle_old
                        for i in range(1, ncols - 1):
                            cur[i] *= scale
                        total = sum(cur)
                    # 若仍超，整体等比收口兜底
                    if total > content_w_mm and total > 0:
                        scale = content_w_mm / total
                        cur = [w * scale for w in cur]
                table.autofit = False
                table.allow_autofit = False
                tblPr = tbl.find(qn('w:tblPr'))
                if tblPr is None:
                    tblPr = OxmlElement('w:tblPr')
                    tbl.insert(0, tblPr)
                layout = tblPr.find(qn('w:tblLayout'))
                if layout is None:
                    layout = OxmlElement('w:tblLayout')
                    tblPr.append(layout)
                layout.set(qn('w:type'), 'fixed')
                self._apply_gridcol_widths(table, cur)
            except Exception:
                pass

        # 文档信息表（sr-copyright-meta）：所有单元格内容水平居中
        # （垂直居中已统一处理，此处补水平居中）
        if is_doc_info_table:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 仅 4.5 节"体检总体评级"表：数据行的第 1、2 列（徽章、风险程度）段落居中
        # 表头行与第 3 列"定义"保持原对齐（与 HTML inline style="text-align:left" 一致）
        if "sr-appendix-grade-tbl" in node_classes:
            for ri, row in enumerate(table.rows):
                if ri == 0:
                    continue
                for ci, cell in enumerate(row.cells):
                    if ci >= 2:
                        continue
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 4.5 节"安全事件/漏洞评级"表：数据行的第 1 列（风险等级）段落居中
        # 与 HTML .sr-appendix-risk-level-tbl td:first-child { text-align:center } 一致
        if "sr-appendix-risk-level-tbl" in node_classes:
            for ri, row in enumerate(table.rows):
                if ri == 0:
                    continue
                for paragraph in row.cells[0].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 表头样式：浅蓝底（#EDF1F7）+ 深色字（#1A1F36），与 HTML .sr-tbl th 保持一致
    _TABLE_HEADER_BG = "EDF1F7"
    _TABLE_HEADER_FG = "1A1F36"

    def _apply_table_header_cell_style(self, cell):
        """给表头单元格应用浅蓝底 + 黑字样式（对齐 HTML .sr-tbl th 视觉）。"""
        self._set_cell_shading(cell, self._TABLE_HEADER_BG)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor.from_string(self._TABLE_HEADER_FG)
                run.bold = True

    def _set_cell_shading(self, cell, hex_color):
        """给单元格设置底色（w:shd 写入 tcPr）。"""
        tcPr = cell._tc.get_or_add_tcPr()
        old = tcPr.find(qn("w:shd"))
        if old is not None:
            tcPr.remove(old)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def _set_cell_no_wrap(self, cell):
        """给单元格设置 <w:noWrap/>，使 Word 不自动换行该单元格内容。

        用于表头和"时间"等需要保持完整不换行的字段。fixed layout 下 Word
        会按 noWrap 单元格的内容宽度自动撑大该列 gridCol。
        """
        tcPr = cell._tc.get_or_add_tcPr()
        old = tcPr.find(qn("w:noWrap"))
        if old is None:
            no_wrap = OxmlElement("w:noWrap")
            tcPr.append(no_wrap)

    def _set_table_column_widths_by_grid(self, table, grid, table_soup=None):
        """根据全表所有行计算列宽并固定到 Word 表格。

        优先级：
          1. 若 table_soup 含 data-col-widths 属性（mm 单位逗号分隔），直接按
             指定宽度设置 gridCol，跳过权重分配。用于人工指定列宽的表（评级表、
             事件表等）。
          2. 否则按权重自动分配 + noWrap 最小列宽兜底（见下文）。

        旧 _set_table_column_widths_by_header 只用 grid[0] 算权重，遇到首行列数
        < max_cols 的表（如文档信息表：首行 2 列、中间行 4 列），zip 会截断，
        剩下列宽度未设导致 Word 按内容撑开溢出。
        本函数扫全 grid，按 colspan 展开到 max_cols 槽位，对每列取该列出现过的
        最大权重，统一归一化到内容宽度。

        策略：
          - 中文字符按 2 宽度权重，英文/数字按 1 宽度
          - 每列权重 = 该列在所有行出现过的单元格权重的最大值
          - 按权重比例分配页面内容宽度（page_width - 2 * margin）
          - 最小列宽 10mm，避免空列宽为 0
          - 关闭 autofit + 设置 tblLayout=fixed，确保 Word 按指定列宽渲染

        Args:
            table: docx Table 对象
            grid: 全表行数据 [[(cell_soup, colspan, rowspan), ...], ...]，
                  与 _map_table 中 grid 一致
        """
        if not grid:
            return
        max_cols = 0
        for row_data in grid:
            cols = sum(int(cs or 1) for _, cs, _ in row_data)
            max_cols = max(max_cols, cols)
        if max_cols == 0:
            return
        cfg = self.config.get("docx", {})
        page_w_mm = cfg.get("page_width_mm", 210)
        margin_mm = cfg.get("margin_mm", 20)
        content_w_mm = page_w_mm - 2 * margin_mm

        # 优先级 1：data-col-widths 人工指定列宽（mm 单位逗号分隔），直接使用
        if table_soup is not None:
            raw_widths = table_soup.get("data-col-widths")
            if raw_widths:
                try:
                    widths = [float(w.strip()) for w in raw_widths.split(",") if w.strip()]
                except (ValueError, TypeError):
                    widths = []
                if len(widths) == max_cols:
                    col_widths_mm = widths
                    # 关闭 autofit，设置 tblLayout=fixed，按指定列宽渲染
                    table.autofit = False
                    table.allow_autofit = False
                    tbl = table._tbl
                    tblPr = tbl.find(qn("w:tblPr"))
                    if tblPr is None:
                        tblPr = OxmlElement("w:tblPr")
                        tbl.insert(0, tblPr)
                    layout = tblPr.find(qn("w:tblLayout"))
                    if layout is None:
                        layout = OxmlElement("w:tblLayout")
                        tblPr.append(layout)
                    layout.set(qn("w:type"), "fixed")
                    self._apply_gridcol_widths(table, col_widths_mm)
                    return
        min_col_mm = 10

        # 扫全 grid，按 colspan 展开到 max_cols 槽位，每列取出现过的最大权重；
        # 同时记录 noWrap 单元格的"需求最小列宽"
        col_weights = [0.0] * max_cols
        col_no_wrap_min = [0.0] * max_cols  # noWrap 单元格需求的最小列宽（mm）
        for row_data in grid:
            ci = 0
            for cell_soup, colspan, rowspan in row_data:
                n = max(int(colspan or 1), 1)
                if cell_soup is None:
                    weight = 0.0
                    text = ""
                else:
                    text = cell_soup.get_text(strip=True) or ""
                    # 中文字符（ord > 127）按 2 宽度，其他按 1
                    weight = sum(2 if ord(c) > 127 else 1 for c in text)
                # 单元格权重按 colspan 平摊到每个槽位
                each_w = weight / n
                for k in range(n):
                    if ci + k < max_cols and each_w > col_weights[ci + k]:
                        col_weights[ci + k] = each_w
                # noWrap 单元格（含 th）按文本宽度计算需求最小列宽
                if cell_soup is not None and (
                    _has_class(cell_soup, "sr-no-wrap") or cell_soup.name == "th"
                ):
                    # 文本宽度估算：中文按 3.7mm/字、ASCII 按 1.9mm/字（10.5pt
                    # 字号下中文实际渲染宽度含字间距），左右各 +2mm padding 共 4mm
                    text_mm = sum(3.7 if ord(c) > 127 else 1.9 for c in text) + 4
                    each_min = text_mm / n
                    for k in range(n):
                        if ci + k < max_cols and each_min > col_no_wrap_min[ci + k]:
                            col_no_wrap_min[ci + k] = each_min
                ci += n
        # 空列兜底权重 1.0，避免后续归一时除 0 或被忽略
        col_weights = [max(w, 1.0) for w in col_weights]

        # 列数过多退化到平均分配；否则按权重比例分配并兜底最小列宽
        if min_col_mm * max_cols >= content_w_mm:
            each = content_w_mm / max_cols
            col_widths_mm = [each] * max_cols
        else:
            total_weight = sum(col_weights)
            col_widths_mm = [content_w_mm * w / total_weight for w in col_weights]
            col_widths_mm = [max(w, min_col_mm) for w in col_widths_mm]
            # 应用 noWrap 需求最小列宽：与权重分配结果取较大值
            col_widths_mm = [max(col_widths_mm[i], col_no_wrap_min[i]) for i in range(max_cols)]
            # 拉底/取较大值后总宽可能超 content_w，按比例归一到 content_w_mm
            total = sum(col_widths_mm)
            if total > content_w_mm and total > 0:
                scale = content_w_mm / total
                col_widths_mm = [w * scale for w in col_widths_mm]

        # 若总宽仍超出 content_w_mm（理论上归一后不会，保留兜底），改用 autofit
        use_autofit = sum(col_widths_mm) > content_w_mm + 0.1

        # 关闭 autofit，设置 tblLayout=fixed
        table.autofit = False
        table.allow_autofit = False
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)

        layout = tblPr.find(qn("w:tblLayout"))
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            tblPr.append(layout)
        layout.set(qn("w:type"), "autofit" if use_autofit else "fixed")

        # tblW 总宽（dxa = twips）
        total_twips = int(round(content_w_mm * 1440 / 25.4))
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.append(tblW)
        tblW.set(qn("w:type"), "dxa")
        tblW.set(qn("w:w"), str(total_twips))

        # 各 gridCol 宽度（twips）
        tblGrid = tbl.find(qn("w:tblGrid"))
        if tblGrid is None:
            return
        for gridCol, width_mm in zip(tblGrid.findall(qn("w:gridCol")), col_widths_mm):
            twips = int(round(width_mm * 1440 / 25.4))
            gridCol.set(qn("w:w"), str(twips))

    def _apply_gridcol_widths(self, table, col_widths_mm):
        """按给定列宽（mm）设置 tblW 总宽 + 各 gridCol 宽度，tblLayout 已由调用方设置。"""
        if not col_widths_mm:
            return
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        total_twips = int(round(sum(col_widths_mm) * 1440 / 25.4))
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.append(tblW)
        tblW.set(qn("w:type"), "dxa")
        tblW.set(qn("w:w"), str(total_twips))
        tblGrid = tbl.find(qn("w:tblGrid"))
        if tblGrid is None:
            return
        for gridCol, width_mm in zip(tblGrid.findall(qn("w:gridCol")), col_widths_mm):
            twips = int(round(width_mm * 1440 / 25.4))
            gridCol.set(qn("w:w"), str(twips))

    def _set_table_column_widths_by_header(self, table, header_row):
        """根据表头单元格文本长度计算列宽并固定到 Word 表格。

        策略：
          - 中文字符按 2 宽度权重，英文/数字按 1 宽度
          - 按权重比例分配页面内容宽度（page_width - 2 * margin）
          - 最小列宽 10mm，避免空表头列宽为 0
          - 关闭 autofit + 设置 tblLayout=fixed，确保 Word 按指定列宽渲染

        Args:
            table: docx Table 对象
            header_row: 表头行数据 [(cell_soup, colspan, rowspan), ...]，
                        与 _map_table 中 grid[0] 一致
        """
        if not header_row:
            return
        cfg = self.config.get("docx", {})
        page_w_mm = cfg.get("page_width_mm", 210)
        margin_mm = cfg.get("margin_mm", 20)
        content_w_mm = page_w_mm - 2 * margin_mm
        min_col_mm = 10

        # 按 colspan 展开，把表头权重铺到每个 gridCol 上
        col_weights = []
        for cell_soup, colspan, rowspan in header_row:
            text = cell_soup.get_text(strip=True) if cell_soup else ""
            # 中文字符（ord > 127）按 2 宽度，其他按 1
            weight = sum(2 if ord(c) > 127 else 1 for c in text)
            n = max(int(colspan or 1), 1)
            col_weights.extend([weight / n] * n)

        if not col_weights:
            return
        col_weights = [max(w, 1.0) for w in col_weights]

        # 列数过多退化到平均分配；否则按权重比例分配并兜底最小列宽
        if min_col_mm * len(col_weights) >= content_w_mm:
            each = content_w_mm / len(col_weights)
            col_widths_mm = [each] * len(col_weights)
        else:
            total_weight = sum(col_weights)
            col_widths_mm = [content_w_mm * w / total_weight for w in col_weights]
            col_widths_mm = [max(w, min_col_mm) for w in col_widths_mm]
            # 拉底后总宽可能超 content_w，按比例归一
            total = sum(col_widths_mm)
            if total > content_w_mm and total > 0:
                scale = content_w_mm / total
                col_widths_mm = [w * scale for w in col_widths_mm]

        # 关闭 autofit，设置 tblLayout=fixed
        table.autofit = False
        table.allow_autofit = False
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)

        layout = tblPr.find(qn("w:tblLayout"))
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            tblPr.append(layout)
        layout.set(qn("w:type"), "fixed")

        # tblW 总宽（dxa = twips）
        total_twips = int(round(content_w_mm * 1440 / 25.4))
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.append(tblW)
        tblW.set(qn("w:type"), "dxa")
        tblW.set(qn("w:w"), str(total_twips))

        # 各 gridCol 宽度（twips）
        tblGrid = tbl.find(qn("w:tblGrid"))
        if tblGrid is None:
            return
        for gridCol, width_mm in zip(tblGrid.findall(qn("w:gridCol")), col_widths_mm):
            twips = int(round(width_mm * 1440 / 25.4))
            gridCol.set(qn("w:w"), str(twips))

    def _render_cell_text_with_br(self, cell, cell_soup):
        """将 HTML 单元格文本写入 docx cell，<br> 和块级标签（<div>、<p>）
        都映射为 Word 软换行。<strong>/<b> → 加粗；<em>/<i> → 斜体。"""
        block_tags = {"div", "p"}
        bold_tags = {"strong", "b"}
        italic_tags = {"em", "i"}
        parts = []
        current_text = []
        children = list(cell_soup.children)
        for i, child in enumerate(children):
            if isinstance(child, Tag) and child.name == "br":
                parts.append(("text", "".join(current_text), None))
                parts.append(("br", None, None))
                current_text = []
            elif isinstance(child, Tag) and child.name in block_tags:
                # 专门处理 sr-component-name-row：把组件名称和类型标签
                # 作为两个独立 run 渲染到同一段落，保留视觉区分
                classes = child.get("class", []) or []
                if isinstance(classes, str):
                    classes = classes.split()
                if "sr-component-name-row" in classes:
                    parts.append(("text", "".join(current_text), None))
                    parts.append(("component_name_row", child, None))
                    current_text = []
                    continue
                if "sr-top5-asset-ip-row" in classes:
                    # Top5 资产行：IP 不加粗 + 托管状态灰色小字括号
                    parts.append(("text", "".join(current_text), None))
                    parts.append(("top5_asset_ip_row", child, None))
                    current_text = []
                    continue
                if "sr-top5-asset-biz" in classes:
                    # 业务资产组名称：在 IP 行下方另起一行（软换行）展示
                    biz_text = child.get_text(" ", strip=True)
                    if biz_text:
                        parts.append(("text", "".join(current_text), None))
                        parts.append(("br", None, None))  # IP 行后换行
                        # 业务名用 9pt 灰色小字（对应 HTML .sr-top5-asset-biz 视觉）
                        parts.append(("text", biz_text, {"color": "6B7A99", "size": 9}))
                        current_text = []
                    continue
                # 块级标签：先把累计文本 flush，再提取块文本
                parts.append(("text", "".join(current_text), None))
                parts.append(("text", child.get_text(" ", strip=True), None))
                # 仅当不是最后一个子元素时才加换行，避免末尾多一个空行
                if i < len(children) - 1:
                    parts.append(("br", None, None))
                current_text = []
            elif isinstance(child, NavigableString):
                current_text.append(str(child))
            elif isinstance(child, Tag) and child.name in bold_tags:
                if current_text:
                    parts.append(("text", "".join(current_text), None))
                    current_text = []
                parts.append(("text", child.get_text(" ", strip=True), "bold"))
            elif isinstance(child, Tag) and child.name in italic_tags:
                if current_text:
                    parts.append(("text", "".join(current_text), None))
                    current_text = []
                parts.append(("text", child.get_text(" ", strip=True), "italic"))
            elif isinstance(child, Tag):
                # 识别 sr-tag--light 系列（浅底深字标签），单独 flush
                child_classes = child.get("class", []) or []
                if isinstance(child_classes, str):
                    child_classes = child_classes.split()
                if any(c == "sr-grade" or c.startswith("sr-grade--") for c in child_classes):
                    # 评级徽章：作为带底色 inline run
                    if current_text:
                        parts.append(("text", "".join(current_text), None))
                        current_text = []
                    parts.append(("grade_badge", child, None))
                elif "sr-tag--light" in child_classes:
                    if current_text:
                        parts.append(("text", "".join(current_text), None))
                        current_text = []
                    parts.append(("tag_light", child, None))
                elif any(c in ("sr-tag--success", "sr-tag--info", "sr-tag--blue",
                               "sr-tag--critical", "sr-tag--high", "sr-tag--medium",
                               "sr-tag--medium-low", "sr-tag--low") for c in child_classes):
                    # 深色 tag 系列：作为带底色 inline run
                    # 4.5 节风险等级表格 cell：按文字（严重/高危/中危/低危）直接映射 HTML 圆点色
                    cell_text = child.get_text(" ", strip=True)
                    if cell_text in ("严重", "高危", "中危", "低危"):
                        if current_text:
                            parts.append(("text", "".join(current_text), None))
                            current_text = []
                        parts.append(("risk_level_text", cell_text, None))
                    else:
                        if current_text:
                            parts.append(("text", "".join(current_text), None))
                            current_text = []
                        parts.append(("tag_dark", child, None))
                else:
                    # 其他 span/容器标签：递归提取，保留嵌套 strong/b/em/i 的 bold/italic 语义
                    self._collect_inline_parts(child, current_text, parts, None)
        if current_text:
            parts.append(("text", "".join(current_text), None))
        # 写入 cell：使用段落+run 的标准方式，<br> 通过 run.add_break() 实现
        paragraph = cell.paragraphs[0]
        # 表格单元格段落段后间距强制为 0（避免继承 Normal 默认 10pt 造成多行表格行高虚高）
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        # 清空默认 run
        for r in list(paragraph.runs):
            r._element.getparent().remove(r._element)
        for ptype, content, fmt in parts:
            if ptype == "text":
                text = (content or "").strip()
                if text:
                    run = paragraph.add_run(text)
                    self._apply_run_fmt(run, fmt)
            elif ptype == "br":
                # 在新 run 中插入软换行（<w:br/>，必须在 <w:r> 内部）
                run = paragraph.add_run()
                run.add_break()
            elif ptype == "component_name_row":
                self._render_component_name_row(paragraph, content)
            elif ptype == "top5_asset_ip_row":
                self._render_top5_asset_ip_row(paragraph, content)
            elif ptype == "grade_badge":
                self._render_grade_badge_inline(paragraph, content)
            elif ptype == "tag_light":
                text = content.get_text(" ", strip=True)
                if text:
                    _light_colors = {
                        "sr-tag--high":       ("FAE8E8", "CF171D"),
                        "sr-tag--warning":    ("FFF1E8", "FA721B"),
                        "sr-tag--success":    ("E7F6F2", "12A679"),
                        "sr-tag--medium":     ("FBF3E7", "D6860D"),
                        "sr-tag--medium-low": ("FBF3E7", "D6860D"),
                        "sr-tag--info":       ("E7F6F8", "0BA7B5"),
                    }
                    tag_classes = content.get("class") or []
                    bg, fg = "EDF1F7", "1A1F36"
                    for cls in tag_classes:
                        if cls in _light_colors:
                            bg, fg = _light_colors[cls]
                            break
                    run = paragraph.add_run(f" {text} ")
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor.from_string(fg)
                    self._set_run_shading(run, bg)
            elif ptype == "tag_dark":
                # 深底白字标签：根据 sr-tag--xxx 取对应色
                classes = content.get("class", []) or []
                if isinstance(classes, str):
                    classes = classes.split()
                tag_key = next((f"tag_{c.replace('sr-tag--', '')}" for c in classes
                                if c.startswith("sr-tag--") and c != "sr-tag--light"), None)
                text = content.get_text(" ", strip=True)
                if text and tag_key:
                    bg, fg = _TAG_COLORS.get(tag_key, ("6F7785", "FFFFFF"))
                    run = paragraph.add_run(f" {text} ")
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(fg)
                    self._set_run_shading(run, bg)
            elif ptype == "risk_level_text":
                # 4.5 节风险等级文字：去底纹，字色=原底纹色，5号(10.5pt)加粗
                text = content
                RISK_LEVEL_COLORS = {
                    "严重": "82010E",   # 原 --sr-critical 底纹色
                    "高危": "CF171D",   # 原 --sr-danger 底纹色
                    "中危": "FDAA1D",   # 原 --sr-caution-dot 底纹色
                    "低危": "6B7A99",   # 原 --sr-text-secondary 底纹色
                }
                if text in RISK_LEVEL_COLORS:
                    fg = RISK_LEVEL_COLORS[text]
                    run = paragraph.add_run(f" {text} ")
                    run.bold = True
                    run.font.size = Pt(10.5)   # 5号字体
                    run.font.color.rgb = RGBColor.from_string(fg)
                    # 不再设置底纹
                else:
                    # 非风险等级文字，按普通文本写入
                    run = paragraph.add_run(text)
                    run.bold = False

    def _render_component_name_row(self, paragraph, div_node):
        """渲染组件名称行：组件名称（默认字号、加粗、黑色）+ 类型（括号形式、灰色小字）
        作为两个独立 run 渲染到同一段落，保留视觉区分度。"""
        # 组件名称
        name_tag = div_node.find("span", class_="sr-component-name")
        if name_tag is not None:
            name_text = name_tag.get_text(strip=True)
            if name_text:
                name_run = paragraph.add_run(name_text)
                name_run.font.size = Pt(10.5)
                name_run.font.bold = True
                name_run.font.color.rgb = RGBColor(0x1F, 0x23, 0x2E)

        # 类型标签（sr-tag sr-tag--light sr-tag--blue）
        type_tag = None
        for span in div_node.find_all("span", class_="sr-tag"):
            classes = span.get("class", []) or []
            if isinstance(classes, str):
                classes = classes.split()
            if "sr-tag--blue" in classes or "sr-tag--light" in classes:
                type_tag = span
                break
        if type_tag is not None:
            type_text = type_tag.get_text(strip=True)
            if type_text:
                # 名称与类型之间留一个空格分隔，类型用括号形式、灰色小字
                paragraph.add_run(" ")
                type_run = paragraph.add_run(f"({type_text})")
                type_run.font.size = Pt(8)
                type_run.font.color.rgb = RGBColor(0x8A, 0x8F, 0x9A)

    def _render_top5_asset_ip_row(self, paragraph, div_node):
        """渲染 Top5 资产行：IP（默认字号、不加粗、黑色）+ 托管状态（括号形式、灰色小字）
        作为两个独立 run 渲染到同一段落，与 _render_component_name_row 风格保持一致。
        IP 不加粗（对应 HTML .sr-top5-asset-ip 的 font-weight:400）。

        HTML 结构：
        <div class="sr-top5-asset-ip-row">
          <span class="sr-top5-asset-ip">10.128.160.30</span>
          <span class="sr-tag sr-tag--light sr-tag--medium">未托管</span>
        </div>
        → Word: "10.128.160.30 (未托管)" 其中 "(未托管)" 是 8pt 灰色小字。"""
        # IP
        ip_tag = div_node.find("span", class_="sr-top5-asset-ip")
        if ip_tag is not None:
            ip_text = ip_tag.get_text(strip=True)
            if ip_text:
                ip_run = paragraph.add_run(ip_text)
                ip_run.font.size = Pt(10.5)
                ip_run.font.bold = False
                ip_run.font.color.rgb = RGBColor(0x1F, 0x23, 0x2E)

        # 托管状态标签（sr-tag--light 系列）
        status_tag = None
        for span in div_node.find_all("span", class_="sr-tag"):
            classes = span.get("class", []) or []
            if isinstance(classes, str):
                classes = classes.split()
            if any(c.startswith("sr-tag--") for c in classes):
                status_tag = span
                break
        if status_tag is not None:
            status_text = status_tag.get_text(strip=True)
            if status_text:
                # IP 与状态之间留一个空格分隔，状态用括号形式、灰色小字
                paragraph.add_run(" ")
                status_run = paragraph.add_run(f"({status_text})")
                status_run.font.size = Pt(8)
                status_run.font.color.rgb = RGBColor(0x8A, 0x8F, 0x9A)

    # ── 工具：shading / borders ─────────────────────

    def _set_table_borders(self, table):
        """给表格加细边框。"""
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{edge}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "4")
            b.set(qn("w:color"), "D3D7DE")
            borders.append(b)
        old = tblPr.find(qn("w:tblBorders"))
        if old is not None:
            tblPr.remove(old)
        tblPr.append(borders)

    def _set_run_shading(self, run, hex_color):
        """给 run 设置底色。"""
        rPr = run._element.get_or_add_rPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        rPr.append(shd)

    # ──────────────────────────────────────────────
    # 编排
    # ──────────────────────────────────────────────

    def run(self):
        """编排：load_config → render_html → snapshot → extract_dom → assemble_docx。"""
        t_total = time.perf_counter()
        try:
            t0 = time.perf_counter()
            self.load_config()
            _log(f"[耗时] load_config: {(time.perf_counter() - t0) * 1000:.0f} ms")

            t0 = time.perf_counter()
            page = self.render_html()
            _log(f"[耗时] render_html: {(time.perf_counter() - t0) * 1000:.0f} ms")

            t0 = time.perf_counter()
            snapshot_map = self.snapshot_complex_components(page)
            _log(f"[耗时] snapshot_complex_components: {(time.perf_counter() - t0) * 1000:.0f} ms ({len(snapshot_map)} 张图)")

            t0 = time.perf_counter()
            root = self.extract_dom(page)
            _log(f"[耗时] extract_dom: {(time.perf_counter() - t0) * 1000:.0f} ms")

            t0 = time.perf_counter()
            self.assemble_docx(root, snapshot_map, self.output_path)
            _log(f"[耗时] assemble_docx: {(time.perf_counter() - t0) * 1000:.0f} ms")

            _log(f"[总耗时] Word 导出完成: {(time.perf_counter() - t_total) * 1000:.0f} ms")
            _log(f"转换完成: {self.output_path}")
        finally:
            self.close()


def main():
    parser = argparse.ArgumentParser(
        description="通用 HTML → Word 转换（混合保真：文本结构化 + 复杂组件截图）"
    )
    parser.add_argument("--input", "-i", default="客户_start_end_安全体检报告.html", help="输入 HTML 文件路径")
    parser.add_argument("--output", "-o", default=None, help="输出 docx 路径（默认 tmp/<输入名>.docx）")
    parser.add_argument("--config", "-c", default=None, help="配置文件 YAML 路径")
    parser.add_argument("--preview-mode", default=None,
                        choices=["a4-landscape", "a4-portrait", "html"],
                        help="覆盖配置中的 preview_mode")
    args = parser.parse_args()
    exporter = HtmlToWordExporter(
        input_path=args.input,
        output_path=args.output,
        config_path=args.config,
        preview_mode=args.preview_mode,
    )
    exporter.run()


if __name__ == "__main__":
    main()
